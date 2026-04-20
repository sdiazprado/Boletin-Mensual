
import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import calendar
import time
import re
from dateutil import parser
import urllib.parse # Necesario para la herramienta de rescate
import cloudscraper  # Para bypass de Cloudflare en BID

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual", layout="wide")

st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    .github-footer {
        position: fixed;
        right: 20px;
        bottom: 20px;
        background-color: rgba(255, 255, 255, 0.9);
        padding: 8px 12px;
        border-radius: 50px;
        border: 1px solid #d0d7de;
        z-index: 1000;
        display: flex;
        align-items: center;
        font-family: 'Calibri', sans-serif;
        text-decoration: none;
        color: #24292f;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.1);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .github-footer:hover {
        transform: translateY(-2px);
        box-shadow: 0px 6px 16px rgba(0,0,0,0.15);
        color: #00205B;
        border-color: #00205B;
    }
    .github-icon {
        margin-right: 8px;
        width: 22px;
        height: 22px;
    }
    </style>
    <a class="github-footer" href="https://github.com/sdiazprado" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@sdiazprado</strong></span>
    </a>
""", unsafe_allow_html=True)

# ==========================================
# UTILIDADES DE FORMATO
# ==========================================
# ==========================================
# HERRAMIENTA DE RESCATE (TEXTO MANUAL)
# ==========================================
@st.cache_data(show_spinner=False)
def buscar_link_inteligente(titulo, organismo):
    """Cazador de DOIs de Doble Impacto (Estricto + Fuzzy). Cero Google."""
    import urllib.parse
    import requests
    import time
    import re

    # 1. Limpieza base
    titulo_raiz = re.split(r'[:\-]', titulo)[0].strip()
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_raiz)
    
    headers = {'User-Agent': 'mailto:bot_investigacion@banco.com'}
    time.sleep(0.5) 

    def consultar_api(query_param, texto_busqueda, modo_estricto=True):
        query_enc = urllib.parse.quote(texto_busqueda)
        url = f"https://api.crossref.org/works?{query_param}={query_enc}&select=URL,title,publisher&rows=4"
        
        try:
            res = requests.get(url, headers=headers, timeout=8)
            if res.status_code == 200:
                items = res.json().get('message', {}).get('items', [])
                
                for item in items:
                    url_oficial = item.get('URL')
                    if not url_oficial: continue
                        
                    pub = item.get('publisher', '').lower()
                    titulo_api = item.get('title', [''])[0].lower()
                    
                    if modo_estricto:
                        if 'oecd' in pub or 'organisation' in pub or organismo.lower() in pub:
                            return url_oficial
                    else:
                        titulo_comparar = titulo_limpio.lower()
                        if titulo_comparar in titulo_api or titulo_api in titulo_comparar:
                            return url_oficial
        except:
            pass
        return None

    link = consultar_api("query.title", titulo_limpio, modo_estricto=True)
    if link: return link

    time.sleep(0.5)
    link = consultar_api("query.bibliographic", titulo, modo_estricto=False)
    if link: return link

    return ""

def procesar_texto_pegado(texto_crudo, organismo_nombre):
    """Extrae Fecha y Título del texto pegado. Retorna DataFrame estandarizado."""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            titulo = ""
            if i >= 1:
                titulo = lineas[i-1]
                basura_menu = ['list view', 'grid view', 'z-a', 'a-z', 'oldest', 'most recent', 'most relevant', 'order by']
                if titulo.lower() in basura_menu and i >= 2: 
                    titulo = lineas[i-2]
            
            if titulo and len(titulo) > 10 and not any(b in titulo.lower() for b in ['search', 'filter', 'sort by', 'publications']):
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo,
                    "Link": "Pendiente",
                    "Organismo": organismo_nombre
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df

def buscar_link_boe(titulo):
    """Busca silenciosamente en la web para obtener el Link Directo y Oficial del BoE"""
    import urllib.parse
    import requests
    from bs4 import BeautifulSoup
    import re
    
    # Extraemos solo el título limpio sin el autor para la búsqueda
    titulo_limpio = titulo.split(': ')[-1] if ': ' in titulo else titulo
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_limpio)
    
    # Usamos DuckDuckGo HTML para evadir bloqueos y obtener el link oficial sin usar Google
    query = f"site:bankofengland.co.uk/speech {titulo_limpio}"
    url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    try:
        res = requests.get(url, headers=headers, timeout=8)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Atrapamos el link real de los resultados
        for a in soup.find_all('a', class_='result__url'):
            href = a.get('href', '').strip()
            if 'bankofengland.co.uk/speech' in href:
                if not href.startswith('http'):
                    href = 'https://' + href
                return href
    except:
        pass
        
    # Fallback de emergencia (1 clic)
    google_query = urllib.parse.quote(query)
    return f"https://www.google.com/search?q={google_query}"

def procesar_texto_pegado_boe(texto_crudo):
    """Extractor especializado para el formato del Bank of England (BoE)"""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            # 1. Buscar Autor un renglón ARRIBA (ej. "Speech // Phil Evans")
            autor = ""
            if i >= 1 and "//" in lineas[i-1]:
                partes = lineas[i-1].split("//")
                if len(partes) > 1:
                    autor = clean_author_name(partes[1].strip())
            
            # 2. Buscar Título Completo dos renglones ABAJO
            titulo = ""
            if i + 2 < len(lineas):
                titulo_raw = lineas[i+2]
                # Le quitamos el sufijo redundante " - speech by Autor"
                titulo_raw = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', titulo_raw).strip()
                titulo = titulo_raw
            
            # 3. Ensamblar y Guardar
            if titulo:
                titulo_final = f"{autor}: {titulo}" if autor else titulo
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo_final,
                    "Link": "Pendiente",
                    "Organismo": "BoE (Inglaterra)"
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df
def clean_author_name(name):
    if not name: return ""
    minusc = ['de', 'van', 'von', 'der', 'del', 'la']
    words = name.strip().split()
    
    # Capitaliza todo excepto las preposiciones europeas
    cleaned_words = [w.capitalize() if w.lower() not in minusc else w.lower() for w in words]
    if cleaned_words:
        cleaned_words[0] = cleaned_words[0].capitalize() # La primera siempre mayúscula
        
    cleaned = " ".join(cleaned_words)
    # Arreglar iniciales pegadas (ej. "J.M. Keynes" -> "J. M. Keynes")
    cleaned = re.sub(r'\b([A-Z])\.\s*([A-Z])', lambda m: f"{m.group(1)}. {m.group(2)}", cleaned)
    return cleaned


# ==========================================
# FUNCIONES DE EXTRACCIÓN (BACKEND)
# ==========================================

# --- SECCIÓN: REPORTES ---
@st.cache_data(show_spinner=False)
def load_reportes_fem(start_date_str, end_date_str):
    """Extractor FEM - Versión Selenium Final (Scroll + Fallback de Fecha)"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://es.weforum.org/publications/"
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")

    try:
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(8)
        # Scroll para despertar la lista dinámica
        driver.execute_script("window.scrollTo(0, 1000);")
        time.sleep(4)
        
        js_script = """
        let res = [];
        document.querySelectorAll('a[href*="/publications/"]').forEach(el => {
            let title = el.innerText || el.textContent || "";
            let container = el.closest('article') || el.closest('div[class*="wef-"]') || el.parentElement;
            let date = container.querySelector('time')?.getAttribute('datetime');
            if (title.length > 15) {
                res.push({ t: title, l: el.href, d: date });
            }
        });
        return res;
        """
        extracted = driver.execute_script(js_script)
        driver.quit()

        for item in extracted:
            # Limpieza de título (quitar saltos de línea y frases de botones)
            titulo = item['t'].split('\n')[0]
            titulo = re.sub(r'(?i)Download PDF|Leer más|Read more|View details', '', titulo).strip()
            link = item['l']
            
            if "/series/" in link: continue

            # Parseo de Fecha
            parsed_date = None
            if item['d']:
                try: parsed_date = parser.parse(item['d']).replace(tzinfo=None)
                except: pass
            
            if not parsed_date:
                # Fallback: Extraer /YYYY/MM/ del link
                m = re.search(r'/(\d{4})/(\d{2})/', link)
                if m: parsed_date = datetime.datetime(int(m.group(1)), int(m.group(2)), 1)

            if parsed_date and start_date <= parsed_date <= end_date:
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FEM"})
    except:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False).drop_duplicates(subset=['Link'])
    return df

# BID (Annual Reports en inglés)
@st.cache_data(show_spinner=False)
def load_reportes_bid_en(start_date_str, end_date_str):
    """
    Extrae Annual Reports del BID en inglés
    URL: https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import pandas as pd
    import time
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango de fechas: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []
    
    # Configuración de paginación
    page = 0
    max_pages = 5  # Límite de páginas a extraer
    hay_resultados = True
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)

    try:
        print("🔍 Iniciando Selenium para BID Annual Reports (EN)...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        while page < max_pages and hay_resultados:
            # URL para Annual Reports en inglés
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports&page={page}"
            
            print(f"\n📄 Accediendo a página {page+1}: {url}")
            driver.get(url)

            try:
                WebDriverWait(driver, 20).until_not(
                    EC.title_contains("Just a moment")
                )
                print(f"✅ Página {page+1} cargada correctamente.")
            except:
                print(f"⚠️ La página {page+1} sigue mostrando 'Just a moment...', esperando...")
                time.sleep(10)

            time.sleep(5)
            html = driver.page_source
            soup = BeautifulSoup(html, 'html.parser')

            # Guardar HTML para depuración (solo primera página)
            if page == 0:
                with open("bid_reportes_debug.html", "w", encoding="utf-8") as f:
                    f.write(html)
                print("💾 HTML guardado en bid_reportes_debug.html")

            # Estrategias de búsqueda
            items = soup.find_all('div', class_='views-row')
            print(f"📚 Página {page+1} - Elementos encontrados: {len(items)}")

            if len(items) == 0:
                print(f"📭 No hay más elementos en página {page+1}")
                hay_resultados = False
                break

            # Mapeo de meses en inglés
            meses_en = {
                'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
            }

            docs_en_pagina = 0
            for idx, item in enumerate(items):
                print(f"\n--- Procesando elemento {idx+1} ---")
                
                # ESTRATEGIA 1: Buscar específicamente el div con clase 'views-field-field-title'
                title_elem = None
                title_container = item.find('div', class_='views-field-field-title')
                if title_container:
                    span_field = title_container.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 1")

                # ESTRATEGIA 2: Buscar span.field-content > a (estructura genérica)
                if not title_elem:
                    span_field = item.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 2")

                # ESTRATEGIA 3: Buscar cualquier enlace con texto largo
                if not title_elem:
                    for a_tag in item.find_all('a', href=True):
                        texto = a_tag.get_text(strip=True)
                        if len(texto) > 30:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 3")
                            break

                if not title_elem:
                    print(f"  ⚠️ No se encontró título en elemento")
                    continue

                titulo = title_elem.get_text(strip=True)
                link = title_elem['href']
                if not link.startswith('http'):
                    link = "https://publications.iadb.org" + link

                print(f"  📌 Título extraído: '{titulo[:100]}...'")

                # Extraer fecha - VERSIÓN MEJORADA
                parsed_date = None
                
                # Buscar específicamente el contenedor de fecha
                date_container = item.find('div', class_='views-field-field-date-issued-text')
                if date_container:
                    date_span = date_container.find('span', class_='field-content')
                    if date_span:
                        date_text = date_span.get_text(strip=True)
                        print(f"  📅 Texto de fecha (específico): {date_text}")
                        
                        # Intentar parsear con regex (ej: "Mar 2026")
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada: {parsed_date}")
                
                # Fallback: buscar cualquier span con texto de fecha
                if not parsed_date:
                    for span in item.find_all('span'):
                        text = span.get_text(strip=True)
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada (fallback): {parsed_date}")
                                break

                if not parsed_date:
                    print(f"  ⚠️ No se pudo extraer fecha")
                    continue

                print(f"  📅 Fecha final: {parsed_date.date()}")

                # Filtrar por fecha
                if parsed_date < start_date or parsed_date > end_date:
                    print(f"  ⏭️ Fecha fuera de rango: {parsed_date.date()} (rango: {start_date.date()} a {end_date.date()})")
                    continue

                # Evitar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "BID (Reportes)"
                    })
                    docs_en_pagina += 1
                    print(f"  ✅ Documento AGREGADO: {titulo[:50]}...")

            print(f"\n📊 Documentos agregados en esta página: {docs_en_pagina}")
            print(f"📊 Total documentos hasta ahora: {len(rows)}")

            page += 1
            print(f"➡️ Avanzando a página {page+1}...\n")

        driver.quit()

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ Documentos BID (Reportes) encontrados en {page} páginas: {len(df)}")
        print("\n📋 Primeros documentos:")
        for i, row in df.head(3).iterrows():
            print(f"  - {row['Date'].strftime('%Y-%m')}: {row['Title'][:80]}...")
    else:
        print("\n⚠️ No se encontraron documentos del BID (Reportes)")

    return df

@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                    except:
                        pass
                if not parsed_date:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div:
                continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag:
                    continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href:
                    continue
                link = "https://www.bis.org" + \
                    href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                    except:
                        pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match:
                        parsed_date = datetime.datetime(
                            int(match.group(1)), 1, 1)
                if not parsed_date:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

# --- Publicaciones Institucionales --- OCDE 

@st.cache_data(show_spinner=False)
def load_pub_inst_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Publicaciones Institucionales (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Pub. Institucionales: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50
    max_pages = 10
    
    print("📡 Solicitando Publicaciones Institucionales a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            # Parámetros para buscar el sub-tema psi114
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-policy-subissues:psi114"  # <-- FILTRO PARA PUB. INSTITUCIONALES
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = data.get("results", [])
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            documentos_en_pagina = 0
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE Pub. Institucionales encontrados: {len(rows)}")
        
    except Exception as e:
        print(f"❌ Error en load_pub_inst_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Pub. Institucionales - Total final: {len(df)}")
    return df



@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """Extractor para Reportes del BM (Solo incluye los que mencionan 'Report')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # ID exacto de la comunidad compartida con Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id,
                'sort': 'dc.date.issued,DESC',
                'page': page,
                'size': 20
            }
            res = requests.get(base_url, headers=headers,
                               params=params, timeout=15)
            data = res.json()

            objects = data.get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects:
                break

            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})

                # Extraer Título y Fecha (Sin Autor, como acordamos)
                title = meta.get('dc.title', [{'value': ''}])[
                    0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[
                    0].get('value', '')

                parsed_date = None
                if date_s:
                    try:
                        parsed_date = parser.parse(date_s)
                    except:
                        pass

                if not parsed_date or parsed_date < start_date:
                    continue

                # --- NUEVO FILTRO PRO-REPORTES ---
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])

                description = ""
                if abstract_list:
                    description = abstract_list[0].get('value', '').lower()
                elif desc_list:
                    description = desc_list[0].get('value', '').lower()

                # Si la palabra "report" NO está en la descripción, lo saltamos
                if not re.search(r'\breport\b', description):
                    continue
                # ----------------------------------

                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[
                    0].get('value', '')
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title,
                                "Link": link, "Organismo": "BM"})
                    items_found += 1

            if items_found == 0:
                break
            page += 1
            if page > 3:
                break  # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"]).dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_reportes_cef(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"https://www.fsb.org/publications/?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all(
                'div', class_=lambda c: c and 'post-excerpt' in c)
            if not items:
                break
            items_found = 0
            for item in items:
                title_div = item.find('div', class_='post-title')
                if not title_div or not title_div.find('a'):
                    continue
                a_tag = title_div.find('a')
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                date_div = item.find('div', class_='post-date')
                parsed_date = None
                if date_div:
                    try:
                        parsed_date = parser.parse(
                            date_div.get_text(strip=True))
                    except:
                        pass
                if not parsed_date:
                    continue
                if not any(r['Link'] == link for r in rows):
                    rows.append(
                        {"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date):
                break
            page += 1
            time.sleep(0.5)
        except:
            break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


# -- OCDE -- REPORTES -- 
@st.cache_data(show_spinner=False)
def load_reportes_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Reports (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []

    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }

    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad
    documentos_procesados = 0

    print("📡 Solicitando Reportes a la API de la OCDE (con paginación)...")

    try:
        while page < max_pages:
            # Parámetros para buscar Reports en inglés
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/reports"  # <-- FILTRO PARA REPORTES
            }

            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)

            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break

            data = response.json()

            # Buscar los resultados
            results = data.get("results", [])

            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break

            documentos_en_pagina = 0
            fecha_mas_antigua = None

            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")

                if not titulo or not link:
                    continue

                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue

                if not parsed_date:
                    continue

                fecha_mas_antigua = parsed_date

                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break

                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()

                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"

                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1

            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")

            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break

            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break

            page += 1
            time.sleep(0.3)  # Pequeña pausa para no sobrecargar la API

        print(f"\n📊 Total Reportes OCDE encontrados: {documentos_procesados}")

    except Exception as e:
        print(f"❌ Error en load_reportes_ocde: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])

    print(f"📊 OCDE Reportes - Total final: {len(df)}")
    return df


@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/bcbspubls.json",
                "https://www.bis.org/api/document_lists/cpmi_publs.json"]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                try:
                    parsed_date = parser.parse(
                        doc.get("publication_start_date", ""))
                except:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except:
            continue
    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div:
                continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag:
                    continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href:
                    continue
                link = "https://www.bis.org" + \
                    href if href.startswith('/') else href
                parsed_date = None
                try:
                    parsed_date = parser.parse(p.get_text(
                        strip=True).replace(titulo, '').strip(', '))
                except:
                    pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match:
                        parsed_date = datetime.datetime(
                            int(match.group(1)), 1, 1)
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except:
            continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# --- SECCIÓN: PUBLICACIONES INSTITUCIONALES ---
def load_pub_inst_oei(start_date_str, end_date_str):
    """Extractor OEI (IEO-IMF) - Versión Selenium con Buscador Recursivo"""
    import undetected_chromedriver as uc
    from bs4 import BeautifulSoup
    import json
    import time
    from dateutil import parser
    import datetime

    # Configuración de fechas
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime.now() - datetime.timedelta(days=365)
        end_date = datetime.datetime.now()

    url = "https://ieo.imf.org/en/publications/annual-reports"
    options = uc.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    # Si usas Linux/Streamlit Cloud, descomenta la siguiente línea:
    # options.binary_location = '/usr/bin/chromium'

    rows = []
    driver = None
    
    try:
        driver = uc.Chrome(options=options)
        driver.get(url)
        time.sleep(6) # Tiempo para hidratación de JS
        
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        script_tag = soup.find('script', id='__NEXT_DATA__')
        
        if script_tag:
            data = json.loads(script_tag.string)
            
            # --- BUSCADOR RECURSIVO ---
            def encontrar_resultados(obj):
                if isinstance(obj, dict):
                    if 'reports' in obj and isinstance(obj['reports'], dict) and 'results' in obj['reports']:
                        return obj['reports']['results']
                    for v in obj.values():
                        res = encontrar_resultados(v)
                        if res: return res
                elif isinstance(obj, list):
                    for item in obj:
                        res = encontrar_resultados(item)
                        if res: return res
                return None

            results = encontrar_resultados(data)
            
            if results:
                for item in results:
                    titulo = item.get('title', {}).get('jsonValue', {}).get('value', '')
                    fecha_raw = item.get('publicationDate', {}).get('jsonValue', {}).get('value', '')
                    
                    # Link (PDF > URL)
                    l_val = item.get('completedReportLink', {}).get('jsonValue', {}).get('value', {})
                    link = l_val.get('href', '') if isinstance(l_val, dict) else ""
                    if not link:
                        link = item.get('url', {}).get('url', '')
                    
                    if link and link.startswith('/'):
                        link = "https://ieo.imf.org" + link
                    
                    if titulo and fecha_raw:
                        p_date = parser.parse(fecha_raw).replace(tzinfo=None)
                        if start_date <= p_date <= end_date:
                            rows.append({
                                "Date": p_date,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": "OEI"
                            })
    except Exception as e:
        print(f"Error Selenium OEI: {e}")
    finally:
        if driver: driver.quit()

    return pd.DataFrame(rows).sort_values("Date", ascending=False)

# ========== FUNCIÓN PARA CEMLA (PUBLICACIONES INSTITUCIONALES) ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Publicaciones Institucionales (Novedades individuales)
    Filtra eventos y contenido no académico
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd
    import time
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*50)
    print("🔍 CEMLA PUBLICACIONES - Extrayendo novedades de boletines...")
    print(f"📅 Rango solicitado: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"✅ Fechas parseadas: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error parseando fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)

    # Palabras a excluir (eventos, cursos, etc. - no publicaciones académicas)
    palabras_excluir = [
        'reunión', 'reunion', 'virtual', 'curso', 'taller', 'seminario',
        'conferencia', 'webinar', 'congreso', 'foro', 'encuentro',
        'junta', 'comité', 'comite', 'próximas actividades', 'calendario',
        'convocatoria', 'premio', 'inscripción', 'registro'
    ]

    rows = []
    
    url = "https://www.cemla.org/comunicados.html"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    }

    try:
        # =========================================================
        # PASO 1: Obtener la lista de boletines
        # =========================================================
        print(f"📡 Solicitando página de boletines: {url}")
        response = requests.get(url, headers=headers, timeout=30, verify=False)
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Error al acceder a la página")
            return pd.DataFrame()

        soup = BeautifulSoup(response.text, 'html.parser')
        
        meses = {
            'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4,
            'mayo': 5, 'junio': 6, 'julio': 7, 'agosto': 8,
            'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
        }
        
        # Encontrar todos los boletines en el rango de fechas
        boletines_a_procesar = []
        
        for ul in soup.find_all('ul', class_='iconlist'):
            for li in ul.find_all('li'):
                p = li.find('p')
                if not p:
                    continue
                
                a_tag = p.find('a')
                if not a_tag:
                    continue
                
                titulo_texto = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                
                match = re.match(r'^([A-Za-z]+)\s+(\d{4})$', titulo_texto, re.IGNORECASE)
                if match:
                    mes_str, año = match.groups()
                    mes_num = meses.get(mes_str.lower(), 0)
                    
                    if mes_num:
                        fecha = datetime.datetime(int(año), mes_num, 1)
                        
                        if start_date <= fecha <= end_date:
                            boletines_a_procesar.append({
                                'fecha': fecha,
                                'titulo': titulo_texto,
                                'link': link
                            })
                            print(f"📌 Boletín encontrado: {fecha.strftime('%Y-%m')} - {titulo_texto}")
        
        print(f"✅ Total boletines en rango: {len(boletines_a_procesar)}")
        
        if not boletines_a_procesar:
            print("⚠️ No se encontraron boletines en el rango de fechas")
            return pd.DataFrame()
        
        # =========================================================
        # PASO 2: Procesar cada boletín y extraer sus novedades
        # =========================================================
        for boletin in boletines_a_procesar:
            print(f"\n🔍 Procesando boletín: {boletin['titulo']} ({boletin['link']})")
            
            try:
                time.sleep(1)
                
                res_boletin = requests.get(boletin['link'], headers=headers, timeout=30, verify=False)
                if res_boletin.status_code != 200:
                    print(f"  ⚠️ Error al acceder al boletín: {res_boletin.status_code}")
                    continue
                
                soup_boletin = BeautifulSoup(res_boletin.text, 'html.parser')
                
                # Buscar todas las novedades (divs con clase "ipost clearfix")
                novedades = soup_boletin.find_all('div', class_=lambda c: c and 'ipost' in c.split() if c else False)
                
                if not novedades:
                    print(f"  ⚠️ No se encontraron novedades en este boletín")
                    continue
                
                print(f"  📚 Novedades encontradas: {len(novedades)}")
                
                for novedad in novedades:
                    try:
                        # Extraer título
                        title_elem = novedad.find('div', class_='entry-title')
                        if not title_elem:
                            continue
                        
                        h3 = title_elem.find('h3')
                        if not h3:
                            continue
                        
                        titulo = h3.get_text(strip=True)
                        
                        # ===== FILTRO: Excluir eventos y contenido no académico =====
                        titulo_lower = titulo.lower()
                        es_excluido = any(palabra in titulo_lower for palabra in palabras_excluir)
                        
                        if es_excluido:
                            print(f"    ⏭️ Excluido (evento): {titulo[:60]}...")
                            continue
                        
                        # Extraer descripción y enlace
                        content_elem = novedad.find('div', class_='entry-content')
                        if not content_elem:
                            continue
                        
                        p = content_elem.find('p')
                        if not p:
                            continue
                        
                        # Extraer el enlace "Leer más..."
                        a_link = p.find('a', href=True)
                        if a_link:
                            link_novedad = a_link.get('href', '')
                            descripcion = p.get_text(strip=True).replace(a_link.get_text(strip=True), '').strip()
                        else:
                            link_novedad = boletin['link']
                            descripcion = p.get_text(strip=True)
                        
                        # Limpiar título
                        titulo = re.sub(r'\s+', ' ', titulo).strip()
                        
                        # Solo agregar si el título es significativo
                        if titulo and len(titulo) > 10:
                            rows.append({
                                'Date': boletin['fecha'],
                                'Title': titulo,
                                'Link': link_novedad if link_novedad else boletin['link'],
                                'Organismo': "CEMLA"
                            })
                            print(f"    ✅ {titulo[:60]}...")
                    
                    except Exception as e:
                        print(f"    ⚠️ Error procesando novedad: {e}")
                        continue
                        
            except Exception as e:
                print(f"  ❌ Error procesando boletín: {e}")
                continue

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'], keep='first')
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ CEMLA PUBLICACIONES - Total novedades: {len(df)} documentos")

    return df


# -- G20 --
@st.cache_data(show_spinner=False)
def load_pub_inst_g20(start_date_str, end_date_str):
    """Extrae documentos del G20 desde la página de News and Media"""
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 G20: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    url = "https://g20.org/media/"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    
    # Palabras clave que queremos incluir
    keywords_incluir = [
        'chair summary', 'declarations', 'g-20 note', 'presidency note',
        'chair\'s summary', 'chair summary', 'g20 note', 'presidency note'
    ]
    
    # Palabras clave para excluir
    keywords_excluir = [
        'agriculture', 'cultura', 'cultural', 'food security', 
        'farming', 'rural', 'agri'
    ]
    
    rows = []

    try:
        print(f"📡 Solicitando página: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"❌ Error al acceder a la página: {res.status_code}")
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Buscar la sección de Press Releases
        press_section = None
        for section in soup.find_all('section', class_='paragraphsection'):
            toptitle = section.find('h2', class_='toptitle')
            if toptitle and 'Press Releases' in toptitle.get_text():
                press_section = section
                break
        
        if not press_section:
            print("⚠️ No se encontró la sección de Press Releases")
            return pd.DataFrame()
        
        # Buscar todos los artículos (h2 seguido de p con fecha)
        articles = press_section.find_all(['h2', 'p'])
        
        i = 0
        while i < len(articles) - 1:
            if articles[i].name == 'h2':
                h2 = articles[i]
                a_tag = h2.find('a')
                
                if a_tag and a_tag.get('href'):
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    
                    if not titulo:
                        i += 1
                        continue
                    
                    if i + 1 < len(articles) and articles[i + 1].name == 'p':
                        p_text = articles[i + 1].get_text(strip=True)
                        
                        match = re.search(r'([A-Za-z]+ \d{1,2},? \d{4})', p_text)
                        if match:
                            fecha_str = match.group(1)
                            try:
                                fecha_str = fecha_str.replace(',', '')
                                parsed_date = datetime.datetime.strptime(fecha_str, '%B %d %Y')
                            except:
                                try:
                                    parsed_date = datetime.datetime.strptime(fecha_str, '%b %d %Y')
                                except:
                                    parsed_date = None
                            
                            if parsed_date:
                                if parsed_date < start_date or parsed_date > end_date:
                                    i += 2
                                    continue
                                
                                titulo_lower = titulo.lower()
                                incluir = any(kw in titulo_lower for kw in keywords_incluir)
                                excluir = any(kw in titulo_lower for kw in keywords_excluir)
                                
                                if excluir or not incluir:
                                    i += 2
                                    continue
                                
                                if link.startswith('/'):
                                    link = f"https://g20.org{link}"
                                
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": "G20"
                                })
                                print(f"   ✅ Agregado: {titulo[:60]}... ({parsed_date.date()})")
                    else:
                        print(f"   ⚠️ No hay párrafo después del h2")
                else:
                    print(f"   ⚠️ h2 sin enlace válido")
            
            i += 1
            
    except Exception as e:
        print(f"❌ Error general: {e}")
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL G20: {len(df)} documentos")
    else:
        print("⚠️ No se encontraron documentos del G20")

    return df

# -- CEF -- #

@st.cache_data(show_spinner=False)
def load_pub_inst_cef(start_date_str, end_date_str):
    url = "https://www.fsb.org/publications/key-regular-publications/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        res = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(res.text, 'html.parser')
        for section in soup.find_all('div', class_='wp-bootstrap-blocks-row'):
            h2 = section.find('h2')
            if not h2:
                continue
            base_title = h2.get_text(strip=True)
            # Latest
            latest_btn = section.find('button', class_='btn-primary')
            if latest_btn and latest_btn.find('a'):
                a_tag = latest_btn.find('a')
                link = "https://www.fsb.org" + \
                    a_tag['href'] if a_tag['href'].startswith(
                        '/') else a_tag['href']
                date_match = re.search(r'\((.*?)\)', a_tag.get_text())
                parsed_date = parser.parse(
                    date_match.group(1)) if date_match else None
                if parsed_date and parsed_date >= start_date:
                    rows.append(
                        {"Date": parsed_date, "Title": f"{base_title}: Latest Report", "Link": link, "Organismo": "CEF"})
            # Previous
            dropdown = section.find('div', class_='dropdown-menu')
            if dropdown:
                for l in dropdown.find_all('a'):
                    year_text = l.get_text(strip=True)
                    try:
                        parsed_date = datetime.datetime(int(year_text), 1, 1)
                    except:
                        parsed_date = None
                    if parsed_date and parsed_date >= start_date:
                        rows.append(
                            {"Date": parsed_date, "Title": f"{base_title} ({year_text})", "Link": l['href'], "Organismo": "CEF"})
    except:
        pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_pub_inst_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/annualeconomicreports.json",
                "https://www.bis.org/api/document_lists/quarterlyreviews.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                try:
                    parsed_date = parser.parse(
                        doc.get("publication_start_date", ""))
                except:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except:
            continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_country_reports_fmi(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Conexión Directa a Coveo API)"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # 1. EL ENDPOINT Y LA LLAVE MAESTRA QUE DESCUBRISTE
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # 2. EL PAYLOAD (Falsificamos la petición del buscador)
    payload = {
        "aq": "@imfseries==\"IMF Staff Country Reports\"",  # Filtro estricto por la Serie
        "numberOfResults": 100,  # Cantidad a traer (Suficiente para un mes)
        "sortCriteria": "@imfdate descending"  # Los más recientes primero
    }

    try:
        # Hacemos un POST directo a la base de datos de Coveo
        res = requests.post(url, headers=headers, json=payload, timeout=15)

        if res.status_code == 200:
            data = res.json()

            # 3. EXTRACCIÓN (Limpia y sin HTML)
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")

                # La fecha viene en timestamp (milisegundos). Lo dividimos entre 1000 para segundos.
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(
                            raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                # Validamos contra la fecha del filtro de la app
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append(
                            {"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_press_releases_fmi(start_date_str, end_date_str):
    """Extractor FMI - Press Releases (Historial completo vía Coveo API)"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # 1. El Endpoint y la llave que tú mismo descubriste
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    # 2. Inyección de Headers para evadir el bloqueo CORS
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",   # <--- LA LLAVE PARA ENTRAR
        "Referer": "https://www.imf.org/",  # <--- CONFIRMA QUE "VENIMOS" DEL FMI
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # 3. Payload: Agregamos el filtro estricto de idioma
    payload = {
        # Le pedimos PRs Y que el idioma sea inglés
        "aq": "@imftype==\"Press Release\" AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)

        if res.status_code == 200:
            data = res.json()

            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")

                # Coveo entrega la fecha en formato Unix (Milisegundos).
                # ¡Es perfecto porque no falla la conversión!
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        # Convertimos de milisegundos a fecha normal
                        parsed_date = datetime.datetime.fromtimestamp(
                            raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                # Filtro final de fechas
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append(
                            {"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_country_reports_elibrary(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Bypass de Tapestry 5 AJAX Lazy-Loading)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []
    base_domain = "https://www.elibrary.imf.org"
    url_overview = f"{base_domain}/view/journals/002/002-overview.xml"

    try:
        # FASE 1: Extraer los tokens dinámicos de AJAX para los años recientes
        res = requests.get(url_overview, headers=headers, timeout=15)
        if res.status_code != 200:
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')

        ajax_links = []
        current_year = datetime.datetime.now().year
        # Buscamos los enlaces de expansión para el año actual y el anterior
        target_years = [str(current_year), str(current_year - 1)]

        for li in soup.find_all('div', attrs={'data-toc-role': 'li'}):
            label_div = li.find('div', class_='label')
            if not label_div:
                continue

            texto_label = label_div.get_text()
            if any(year in texto_label for year in target_years):
                a_tag = li.find('a', class_='ajax-control')
                if a_tag and a_tag.has_attr('href'):
                    ajax_links.append(base_domain + a_tag['href'])

        # FASE 2: Interceptar y "deshidratar" las respuestas AJAX de Tapestry
        headers_ajax = headers.copy()
        # Engañamos al framework
        headers_ajax['X-Requested-With'] = 'XMLHttpRequest'
        headers_ajax['Accept'] = 'application/json, text/javascript, */*; q=0.01'

        for ajax_url in ajax_links:
            try:
                res_ajax = requests.get(
                    ajax_url, headers=headers_ajax, timeout=15)
                if res_ajax.status_code != 200:
                    continue

                data = res_ajax.json()

                # Extraemos el HTML inyectado dentro del nodo "zones"
                html_fragment = ""
                if "zones" in data:
                    for zone_id, html_content in data["zones"].items():
                        html_fragment += html_content

                if not html_fragment:
                    continue

                # FASE 3: Parsear el HTML revelado
                soup_fragment = BeautifulSoup(html_fragment, 'html.parser')

                for a_tag in soup_fragment.find_all('a', href=True):
                    href = a_tag['href']
                    titulo = a_tag.get_text(strip=True)

                    # Filtro de sanidad: debe ser un artículo real
                    if '/view/journals/002/' in href and len(titulo) > 15:
                        link_real = base_domain + \
                            href if href.startswith('/') else href

                        # Buscamos la fecha subiendo hasta 3 niveles en el DOM
                        date_str = ""
                        for padre in a_tag.find_parents(['div', 'li'], limit=3):
                            texto_padre = padre.get_text(
                                separator=" ", strip=True)

                            # Caza fechas en formatos "Mar 05, 2026" o "05 March 2026"
                            match = re.search(
                                r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}', texto_padre)
                            if not match:
                                match = re.search(
                                    r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}', texto_padre)

                            if match:
                                date_str = match.group(0)
                                break  # Encontramos la fecha, salimos del bucle

                        parsed_date = None
                        if date_str:
                            try:
                                parsed_date = parser.parse(date_str)
                                if parsed_date.tzinfo is not None:
                                    parsed_date = parsed_date.replace(
                                        tzinfo=None)
                            except:
                                pass

                        # Evaluación final
                        if parsed_date and parsed_date >= start_date:
                            if not any(r['Link'] == link_real for r in rows):
                                rows.append(
                                    {"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
            except:
                continue  # Aislamiento de fallos

    except Exception as e:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## FMI - Publiccaciones Institucionales - INICIO

## FMI - F&D Magazine (inicio)
@st.cache_data(show_spinner=False)
def load_pub_inst_fandd(start_date_str, end_date_str):
    """Extrae ediciones completas de la revista F&D Magazine del FMI"""
    import requests
    import json
    import re
    import datetime
    import pandas as pd

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI F&D: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    url = "https://www.imf.org/en/publications/fandd/issues"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

    rows = []

    try:
        print(f"📡 Solicitando página: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"❌ Error al acceder a la página: {res.status_code}")
            return pd.DataFrame()

        match = re.search(r'<script id="__NEXT_DATA__" type="application/json">(.*?)</script>', res.text, re.DOTALL)
        if not match:
            print("❌ No se encontró el script __NEXT_DATA__")
            return pd.DataFrame()

        data = json.loads(match.group(1))
        
        # Navegación por componentProps
        try:
            component_props = data['props']['pageProps']['componentProps']
            issue_list = None
            for comp_id, comp_data in component_props.items():
                if 'issueList' in comp_data:
                    issue_list = comp_data['issueList']
                    print(f"✅ Encontrado issueList en componente: {comp_id}")
                    break
            
            if not issue_list:
                print("❌ No se encontró issueList en componentProps")
                return pd.DataFrame()
            
            results = issue_list.get('results', [])
            print(f"✅ Total de números encontrados: {len(results)}")
            
        except (KeyError, TypeError) as e:
            print(f"❌ Error navegando: {e}")
            return pd.DataFrame()

        meses_map = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }

        for issue in results:
            issue_title = issue.get('issueTitle', {}).get('jsonValue', {}).get('value', '').strip()
            issue_label = issue.get('issueLabel', {}).get('jsonValue', {}).get('value', '').strip()
            issue_url = issue.get('url', {}).get('url', '')
            
            fecha_texto = issue_label if issue_label else issue_title
            
            match_date = re.search(r'([A-Za-z]+)\s+(\d{4})', fecha_texto, re.IGNORECASE)
            if not match_date:
                print(f"   ⚠️ No se pudo parsear fecha de: '{fecha_texto}'")
                continue
            
            mes_str = match_date.group(1).lower()
            año = int(match_date.group(2))
            mes_num = meses_map.get(mes_str, 1)
            issue_date = datetime.datetime(año, mes_num, 19)
            
            if issue_date < start_date or issue_date > end_date:
                continue
            
            title_clean = re.sub(r'\s+', ' ', issue_title).strip()
            if not title_clean:
                title_clean = fecha_texto
            
            rows.append({
                "Date": issue_date,
                "Title": title_clean,
                "Link": issue_url,
                "Organismo": "F&D Magazine"
            })
        
    except Exception as e:
        print(f"❌ Error general: {e}")
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL FMI F&D: {len(df)} ediciones")

    return df

## FMI - 

@st.cache_data(show_spinner=False)
def load_pub_inst_fmi(start_date_str, end_date_str):
    """Extractor FMI - Vía directa por API Next.js (El Regalo)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*'
    }

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # 1. CAZADOR DE BUILD ID (Para que tu código no caduque nunca)
    build_id = "OPXKbpp2La91iW-gTVkBX"  # Tu regalo como plan de respaldo
    try:
        res_html = requests.get(
            "https://www.imf.org/en/publications", headers=headers, timeout=15)
        # Buscamos el código dinámico oculto en la página principal
        match = re.search(r'"buildId":"([^"]+)"', res_html.text)
        if match:
            build_id = match.group(1)
    except:
        pass

    # 2. CONSTRUCCIÓN DE LOS ENLACES JSON DIRECTOS
    endpoints_json = [
        f"https://www.imf.org/_next/data/{build_id}/en/publications/fm.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/weo.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/gfsr.json"
    ]

    for url in endpoints_json:
        try:
            # Ahora pedimos el JSON limpio, evadiendo el HTML
            res = requests.get(url, headers=headers, timeout=15)
            if res.status_code != 200:
                continue
            data = res.json()

            # Buscador recursivo dentro del JSON
            def extraer_issues(obj):
                if isinstance(obj, dict):
                    if "issuePage" in obj and isinstance(obj["issuePage"], dict) and "results" in obj["issuePage"]:
                        for r in obj["issuePage"]["results"]:
                            yield r
                    for k, v in obj.items():
                        yield from extraer_issues(v)
                elif isinstance(obj, list):
                    for item in obj:
                        yield from extraer_issues(item)

            for issue in extraer_issues(data):
                titulo = issue.get("title", {}).get(
                    "jsonValue", {}).get("value", "")
                link_raw = issue.get("url", {}).get(
                    "url", "") or issue.get("url", {}).get("path", "")
                if not titulo or not link_raw:
                    continue

                link_real = link_raw if link_raw.startswith(
                    "http") else "https://www.imf.org" + link_raw

                d_str = issue.get("publicationDate", {}).get(
                    "jsonValue", {}).get("value", "")
                if d_str:
                    try:
                        parsed_date = parser.parse(d_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                        if parsed_date >= start_date and not any(r['Link'] == link_real for r in rows):
                            rows.append(
                                {"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
                    except:
                        pass
        except:
            continue

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_pub_inst_bm(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales (Colecciones Específicas) del BM"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # IDs exactos de las 3 colecciones
    scopes = [
        '4c48a649-7773-4d0f-b441-f5fc7e8d67f8',  # Business Ready
        '09c5e8fc-187f-5c2f-a077-3e03044c7b62',  # Perspectivas económicas mundiales
        '3d9bbbf6-c007-5043-b655-04d8a1cfbfb2'  # Tercera colección
    ]

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # Iteramos sobre cada una de las colecciones
    for scope in scopes:
        page = 0
        while True:
            try:
                # Al pasarle el 'scope', la API restringe la búsqueda SOLO a esa colección
                params = {
                    'scope': scope,
                    'sort': 'dc.date.issued,DESC',
                    'page': page,
                    'size': 20
                }
                res = requests.get(base_url, headers=headers,
                                   params=params, timeout=15)
                data = res.json()

                objects = data.get('_embedded', {}).get(
                    'searchResult', {}).get('_embedded', {}).get('objects', [])
                if not objects:
                    break

                items_found = 0
                for obj in objects:
                    item = obj.get('_embedded', {}).get('indexableObject', {})
                    meta = item.get('metadata', {})

                    title = meta.get('dc.title', [{'value': ''}])[
                        0].get('value', '')
                    date_s = meta.get('dc.date.issued', [{'value': ''}])[
                        0].get('value', '')

                    parsed_date = None
                    if date_s:
                        try:
                            parsed_date = parser.parse(date_s)
                        except:
                            pass

                    if not parsed_date or parsed_date < start_date:
                        continue

                    link = meta.get('dc.identifier.uri', [{'value': ''}])[
                        0].get('value', '')
                    if not link:
                        link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                    if not any(r['Link'] == link for r in rows):
                        rows.append(
                            {"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                        items_found += 1

                if items_found == 0:
                    break
                page += 1
                if page > 3:
                    break  # Límite de seguridad
                time.sleep(0.2)
            except:
                break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

    # --- SECCIÓN: INVESTIGACIÓN ---
@st.cache_data(show_spinner=False)
def load_working_papers_fmi(start_date_str, end_date_str):
    """Extractor FMI - Working Papers (Conexión Directa Coveo API mediante @imfseries)"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # Búsqueda exacta por la serie "IMF Working Papers"
    payload = {
        "aq": "@imfseries==\"IMF Working Papers\" AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)

        if res.status_code == 200:
            data = res.json()

            for item in data.get("results", []):
                titulo = item.get("title", "").strip()
                link = item.get("clickUri", "")

                # Extraer Fecha (De milisegundos Unix a Datetime normal)
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                # Filtrar por el rango de fechas de la app
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date, 
                            "Title": titulo, 
                            "Link": link, 
                            "Organismo": "FMI"
                        })
    except Exception as e:
        pass # Silencioso para no romper la app principal

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


# ========== FUNCIÓN PARA CEMLA (PUBLICACIONES INSTITUCIONALES) ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_cemla(start_date_str, end_date_str):
    """Extractor para Boletín CEMLA - Versión optimizada para la estructura de Mailchimp"""
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd
    import time

    url = "https://www.cemla.org/comunicados.html"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

    print("="*50)
    print("🔍 Iniciando extracción de CEMLA con novedades individuales...")
    print(f"📅 Rango solicitado: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"✅ Fechas parseadas: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error parseando fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)
        print(f"📅 Usando rango por defecto: {start_date.date()} a {end_date.date()}")

    rows = []

    # Palabras y URLs a excluir
    palabras_excluir = [
        'convocatoria', 'premio', 'curso', 'taller', 'seminario',
        'evento', 'webinar', 'congreso', 'beca', 'inscripción',
        'registro', 'participación', 'invitación', 'calendario',
        'programa de actividades', 'agenda', 'convocan', 'postulación',
        'reunión de gobernadores', 'reunión de responsables', 'encuesta',
        'award', 'prize', 'conference', 'workshop', 'registration',
        'call for papers', 'agenda', 'calendar', 'program', 'invitation',
        'survey', 'meeting', 'governors', 'responsables'
    ]
    urls_excluir = [
        'calendario', 'premiodebancacentral', 'convocatoria',
        'award', 'prize', 'course', 'workshop', 'event',
        'reunion', 'meeting', 'programa-actividades'
    ]

    meses_map = {
        'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4,
        'mayo': 5, 'junio': 6, 'julio': 7, 'agosto': 8,
        'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12,
    }

    try:
        print(f"📡 Solicitando lista de boletines desde {url}...")
        res = requests.get(url, headers=headers, timeout=15)
        print(f"   Status code: {res.status_code}")
        
        if res.status_code != 200:
            print(f"   ❌ Error al acceder a la página")
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')
        print(f"✅ Página cargada, {len(soup.text)} caracteres")

        # ===== 1. EXTRAER LISTA DE BOLETINES =====
        boletines = []
        for element in soup.find_all(['p', 'div', 'h3', 'h4', 'li']):
            text = element.get_text(strip=True)
            match = re.match(r'^([A-Za-z]+)\s+(\d{4})', text)
            if not match:
                continue

            mes_str, year_str = match.groups()
            mes_num = meses_map.get(mes_str.lower())
            if not mes_num:
                print(f"   ⚠️ Mes no reconocido: {mes_str}")
                continue

            try:
                fecha = datetime.datetime(int(year_str), mes_num, 1)
            except Exception as e:
                print(f"   ⚠️ Error fecha: {e}")
                continue

            a_tag = element.find('a', href=True, string=re.compile(r'Ver más', re.I))
            if not a_tag:
                next_elem = element.find_next_sibling()
                if next_elem:
                    a_tag = next_elem.find('a', href=True, string=re.compile(r'Ver más', re.I))
            
            if a_tag:
                href = a_tag.get('href')
                if href:
                    if href.startswith('/'):
                        link = f"https://www.cemla.org{href}"
                    elif href.startswith('http'):
                        link = href
                    else:
                        link = f"https://www.cemla.org/{href}"
                    
                    boletines.append({
                        'fecha': fecha,
                        'titulo': text,
                        'link': link
                    })
                    print(f"📌 Boletín encontrado: {fecha.strftime('%Y-%m')} - {text[:50]}...")

        print(f"✅ Total boletines principales: {len(boletines)}")

        if not boletines:
            print("⚠️ No se encontraron boletines. Verifica la estructura de la página.")
            with open("cemla_debug.html", "w", encoding="utf-8") as f:
                f.write(res.text)
            print("💾 HTML guardado en cemla_debug.html para depuración")
            return pd.DataFrame()

        # ===== 2. PROCESAR CADA BOLETÍN =====
        for boletin in boletines:
            if boletin['fecha'] < start_date or boletin['fecha'] > end_date:
                print(f"⏭️ Boletín fuera de rango: {boletin['fecha'].strftime('%Y-%m')}")
                continue

            print(f"\n🔍 Procesando boletín {boletin['fecha'].strftime('%Y-%m')}: {boletin['link']}")
            
            try:
                time.sleep(1)
                
                res_boletin = requests.get(boletin['link'], headers=headers, timeout=15)
                if res_boletin.status_code != 200:
                    print(f"  ⚠️ Error al acceder al boletín: {res_boletin.status_code}")
                    continue

                soup_boletin = BeautifulSoup(res_boletin.text, 'html.parser')
                
                novedades = []
                
                # ===== ESTRATEGIA MEJORADA: Buscar bloques de novedades =====
                # Busca divs con clase "ipost clearfix" o similar (estructura de Mailchimp)
                bloques = soup_boletin.find_all('div', class_=lambda c: c and 'ipost' in c.split())
                
                if not bloques:
                    # Fallback: buscar cualquier div que contenga un h3 y un enlace
                    bloques = soup_boletin.find_all('div', class_=lambda c: c and ('entry' in c or 'post' in c))
                
                print(f"   Bloques de novedades encontrados: {len(bloques)}")
                
                for bloque in bloques:
                    try:
                        # 1. Extraer título del bloque (desde h3)
                        title_elem = bloque.find('h3')
                        if not title_elem:
                            title_elem = bloque.find(['h1', 'h2', 'h4'])
                        
                        if not title_elem:
                            continue
                        
                        titulo = title_elem.get_text(strip=True)
                        if not titulo or len(titulo) < 10:
                            continue
                        
                        # 2. Buscar enlace relevante dentro del bloque
                        link_final = None
                        enlaces = bloque.find_all('a', href=True)
                        
                        for a in enlaces:
                            href = a.get('href', '').strip()
                            if not href:
                                continue
                            
                            # Excluir enlaces de redes sociales, suscripción, etc.
                            if any(x in href.lower() for x in ['twitter', 'facebook', 'mailchi.mp', 'unsubscribe', 'share', 'forward']):
                                continue
                            
                            # Construir URL absoluta
                            if href.startswith('/'):
                                href_full = f"https://www.cemla.org{href}"
                            elif href.startswith('http'):
                                href_full = href
                            else:
                                href_full = f"https://www.cemla.org/{href}"
                            
                            # Priorizar PDFs o enlaces que no sean "Leer más"
                            if href_full.endswith('.pdf') or not re.search(r'leer\s*más', a.get_text(strip=True), re.I):
                                link_final = href_full
                                break
                            else:
                                if not link_final:
                                    link_final = href_full
                        
                        if not link_final:
                            continue
                        
                        # 3. Limpiar título
                        titulo = re.sub(r'\s+', ' ', titulo).strip()
                        if len(titulo) > 150:
                            titulo = titulo[:150] + "..."
                        
                        # 4. Verificar exclusión
                        texto_lower = titulo.lower()
                        url_lower = link_final.lower()
                        
                        es_excluido_titulo = any(p in texto_lower for p in palabras_excluir)
                        es_excluido_url = any(p in url_lower for p in urls_excluir)
                        
                        if es_excluido_titulo or es_excluido_url:
                            print(f"  ⏭️ Excluido: {titulo[:50]}...")
                            continue
                        
                        # 5. Agregar a novedades
                        novedades.append({
                            'Date': boletin['fecha'],
                            'Title': titulo,
                            'Link': link_final,
                            'Organismo': "CEMLA"
                        })
                        print(f"  ✅ {titulo[:60]}...")
                        
                    except Exception as e:
                        print(f"    ❌ Error procesando bloque: {e}")
                        continue
                
                if not novedades:
                    print("  ⚠️ No se encontraron enlaces relevantes. Primeros 5 enlaces del boletín:")
                    for i, a in enumerate(soup_boletin.find_all('a', href=True)[:5]):
                        href = a.get('href', '')
                        texto = a.get_text(strip=True) or "SIN TEXTO"
                        print(f"     {i+1}. Texto: '{texto[:60]}' -> URL: {href[:80]}")
                
                rows.extend(novedades)
                print(f"  📊 Total novedades en este boletín: {len(novedades)}")
                
            except Exception as e:
                print(f"  ❌ Error procesando boletín: {e}")
                continue

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        
        print(f"\n🔍 Eliminando duplicados...")
        df = df.drop_duplicates(subset=['Date', 'Link'], keep='first')
        
        enlaces_a_excluir = [
            'twitter.com/share',
            'mailchi.mp/cemla.org/boletin',
            'e=UNIQID'
        ]
        for excluir in enlaces_a_excluir:
            df = df[~df['Link'].str.contains(excluir, na=False)]
        
        print(f"   Después: {len(df)} registros")
        df = df.sort_values("Date", ascending=False)

        print(f"\n✅ TOTAL CEMLA PUBLICACIONES: {len(df)} documentos")
        if not df.empty:
            print("📋 PRIMEROS 3 DOCUMENTOS:")
            for i, row in df.head(3).iterrows():
                print(f"   - {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:60]}...")
    else:
        print("⚠️ No se encontraron novedades")

    return df

## INVESTIGACIÓN 
@st.cache_data(show_spinner=False)
def load_investigacion_bpi(start_date_str, end_date_str):
    """Extractor Investigación BPI (BIS Papers & Working Papers) vía API JSON"""
    import requests
    import pandas as pd
    import datetime
    import html
    from dateutil import parser
    
    # Configuración de fechas
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime.now() - datetime.timedelta(days=365)
        end_date = datetime.datetime.now()

    # Los dos endpoints JSON de investigación del BPI
    urls = [
        "https://www.bis.org/api/document_lists/bispapers.json",
        "https://www.bis.org/api/document_lists/wppubls.json"
    ]
    
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                data = response.json()
                
                # Usamos la lógica de extracción exacta de tu función de discursos
                for path, doc in data.get("list", {}).items():
                    title = html.unescape(doc.get("short_title", ""))
                    date_str = doc.get("publication_start_date", "")
                    
                    link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                    
                    if title and date_str:
                        try:
                            # Convertimos a datetime para aplicar el filtro de tu app
                            p_date = parser.parse(date_str).replace(tzinfo=None)
                            
                            if start_date <= p_date <= end_date:
                                rows.append({
                                    "Date": p_date,
                                    "Title": title,
                                    "Link": link,
                                    "Organismo": "BPI"
                                })
                        except:
                            continue
        except Exception as e:
            print(f"Error BPI Investigación en {url}: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        
    return df

## BID - Inglés 
@st.cache_data(show_spinner=False)
def load_investigacion_bid_en(start_date_str, end_date_str):
    """
    Extrae Working Papers del BID en inglés usando undetected-chromedriver
    """
    import undetected_chromedriver as uc
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    import ssl
    import urllib3
    
    # 🔧 SOLUCIÓN PARA REDES CORPORATIVAS
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    ssl._create_default_https_context = ssl._create_unverified_context
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Inglés: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    page = 0
    max_pages = 3
    
    options = uc.ChromeOptions()
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_argument('--disable-gpu')
    
    meses_en = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    try:
        print("🔍 Iniciando BID Inglés...")
        driver = uc.Chrome(options=options, version_main=146)
        time.sleep(2)
        
        while page < max_pages:
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers&page={page}"
            print(f"📄 Página {page+1}: {url}")
            
            driver.get(url)
            time.sleep(10)
            
            if "Just a moment" in driver.page_source:
                print("   ⚠️ Cloudflare detectado, esperando...")
                time.sleep(15)
            
            try:
                WebDriverWait(driver, 45).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "views-row"))
                )
            except:
                pass
            
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay artículos en página {page+1}")
                break
            
            print(f"   📚 Artículos: {len(items)}")
            
            for item in items:
                try:
                    title_container = item.find('div', class_='views-field-field-title')
                    if title_container:
                        a_tag = title_container.find('a')
                        if a_tag:
                            titulo = a_tag.get_text(strip=True)
                            link = a_tag.get('href', '')
                            if link and not link.startswith('http'):
                                link = "https://publications.iadb.org" + link
                    else:
                        continue
                    
                    if not titulo or len(titulo) < 10:
                        continue
                    
                    date_container = item.find('div', class_='views-field-field-date-issued-text')
                    if date_container:
                        date_text = date_container.get_text(strip=True)
                        match = re.search(r'([A-Za-z]+)\s+(\d{4})', date_text)
                        if match:
                            mes_str = match.group(1).lower()
                            año = int(match.group(2))
                            mes_num = meses_en.get(mes_str, 1)
                            parsed_date = datetime.datetime(año, mes_num, 1)
                        else:
                            continue
                    else:
                        continue
                    
                    # Filtrar por año y mes
                    if parsed_date.year < start_date.year or parsed_date.year > end_date.year:
                        continue
                    if parsed_date.year == start_date.year and parsed_date.month < start_date.month:
                        continue
                    if parsed_date.year == end_date.year and parsed_date.month > end_date.month:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Inglés)"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
                        
                except Exception as e:
                    print(f"   ⚠️ Error procesando artículo: {e}")
                    continue
            
            page += 1
            time.sleep(3)
        
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error en BID Inglés: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ BID Inglés: {len(df)} documentos")
    else:
        print("\n⚠️ No se encontraron documentos del BID (Inglés)")
    
    return df

## BID ESPAÑOL 

@st.cache_data(show_spinner=False)
def load_investigacion_bid(start_date_str, end_date_str):
    import undetected_chromedriver as uc
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    import ssl
    import urllib3
    
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    ssl._create_default_https_context = ssl._create_unverified_context
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Español: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    page = 0
    max_pages = 3
    
    options = uc.ChromeOptions()
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_argument('--disable-gpu')
    
    meses_es = {
        'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
        'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12,
    }
    
    try:
        print("🔍 Iniciando BID Español...")
        driver = uc.Chrome(options=options, version_main=146)
        time.sleep(2)
        
        while page < max_pages:
            url = f"https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo&page={page}"
            print(f"📄 Página {page+1}: {url}")
            
            driver.get(url)
            time.sleep(10)
            
            if "Just a moment" in driver.page_source:
                print("   ⚠️ Cloudflare detectado, esperando...")
                time.sleep(15)
            
            try:
                WebDriverWait(driver, 45).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "views-row"))
                )
            except:
                pass
            
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay artículos en página {page+1}")
                break
            
            print(f"   📚 Artículos: {len(items)}")
            
            for item in items:
                try:
                    title_container = item.find('div', class_='views-field-field-title')
                    if title_container:
                        a_tag = title_container.find('a')
                        if a_tag:
                            titulo = a_tag.get_text(strip=True)
                            link = a_tag.get('href', '')
                            if link and not link.startswith('http'):
                                link = "https://publications.iadb.org" + link
                    else:
                        continue
                    
                    if not titulo or len(titulo) < 10:
                        continue
                    
                    date_container = item.find('div', class_='views-field-field-date-issued-text')
                    if date_container:
                        date_text = date_container.get_text(strip=True)
                        match = re.search(r'([A-Za-z]+)\s+(\d{4})', date_text)
                        if match:
                            mes_str = match.group(1).lower()
                            año = int(match.group(2))
                            mes_num = meses_es.get(mes_str, 1)
                            parsed_date = datetime.datetime(año, mes_num, 1)
                        else:
                            continue
                    else:
                        continue
                    
                    if parsed_date.year < start_date.year or parsed_date.year > end_date.year:
                        continue
                    if parsed_date.year == start_date.year and parsed_date.month < start_date.month:
                        continue
                    if parsed_date.year == end_date.year and parsed_date.month > end_date.month:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
                        
                except Exception as e:
                    continue
            
            page += 1
            time.sleep(3)
        
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ BID Español: {len(df)} documentos")
    
    return df
## CEMLA INVESTIGACIÓN - 
@st.cache_data(show_spinner=False)
def load_investigacion_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Latin American Journal of Central Banking
    Extrae fecha COMPLETA (con día) si está disponible en Crossref
    """
    import requests
    import datetime
    from dateutil import parser
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*60)
    print("🔍 CEMLA INVESTIGACIÓN - Buscando fechas completas (con día)")
    print("="*60)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    issn = "2666-1438"
    base_url = "https://api.crossref.org/works"
    
    # Buscar mes por mes para tener mejor control
    current = start_date.replace(day=1)
    
    while current <= end_date:
        year = current.year
        month = current.month
        
        # Último día del mes
        if month == 12:
            last_day = 31
        elif month in [4, 6, 9, 11]:
            last_day = 30
        else:
            last_day = 28 if year % 4 != 0 else 29
        
        fecha_inicio = f"{year}-{month:02d}-01"
        fecha_fin = f"{year}-{month:02d}-{last_day}"
        
        print(f"\n📆 Buscando {year}-{month:02d}...")
        
        params = {
            "filter": f"from-pub-date:{fecha_inicio},until-pub-date:{fecha_fin},issn:{issn}",
            "rows": 50,
            "sort": "published-online",
            "order": "desc"
        }
        
        try:
            response = requests.get(base_url, params=params, timeout=30, verify=False)
            
            if response.status_code == 200:
                data = response.json()
                items = data.get('message', {}).get('items', [])
                
                if items:
                    print(f"   📚 Artículos: {len(items)}")
                    
                    for item in items:
                        titulo = item.get('title', [''])[0] if item.get('title') else ''
                        doi = item.get('DOI', '')
                        link = f"https://doi.org/{doi}" if doi else ''
                        
                        if not titulo or not link:
                            continue
                        
                        # ========== INTENTAR OBTENER FECHA COMPLETA ==========
                        fecha_completa = None
                        
                        # 1. Probar con 'published-online' (puede tener día)
                        pub_online = item.get('published-online', {})
                        if pub_online:
                            date_parts = pub_online.get('date-parts', [[]])[0]
                            if len(date_parts) >= 3:
                                try:
                                    fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                    print(f"      📅 Online: {fecha_completa.strftime('%Y-%m-%d')}")
                                except:
                                    pass
                        
                        # 2. Probar con 'issued' (fecha de publicación)
                        if not fecha_completa:
                            issued = item.get('issued', {})
                            if issued:
                                date_parts = issued.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Issued: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 3. Probar con 'posted-online'
                        if not fecha_completa:
                            posted = item.get('posted-online', {})
                            if posted:
                                date_parts = posted.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Posted: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 4. Fallback: usar el primer día del mes (si no hay día)
                        if not fecha_completa:
                            fecha_completa = datetime.datetime(year, month, 1)
                            print(f"      ⚠️ Fallback: {fecha_completa.strftime('%Y-%m-%d')} (sin día específico)")
                        
                        # Filtrar por rango
                        if start_date <= fecha_completa <= end_date:
                            rows.append({
                                "Date": fecha_completa,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": "CEMLA"
                            })
                            print(f"      ✅ AGREGADO: {fecha_completa.strftime('%Y-%m-%d')}")
                        else:
                            print(f"      ⏭️ Fuera de rango: {fecha_completa.strftime('%Y-%m-%d')}")
                            
                else:
                    print(f"   📭 Sin artículos")
                    
            else:
                print(f"   ❌ Error: {response.status_code}")
                
        except Exception as e:
            print(f"   ❌ Error: {e}")
        
        # Siguiente mes
        if current.month == 12:
            current = current.replace(year=current.year + 1, month=1)
        else:
            current = current.replace(month=current.month + 1)
        
        time.sleep(0.5)
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n{'='*60}")
    print(f"📊 CEMLA Investigación - Total: {len(df)} documentos")
    if not df.empty:
        print("\n📅 Primeros 5 documentos con sus fechas:")
        for i, row in df.head(5).iterrows():
            print(f"   {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:60]}...")
    print(f"{'='*60}")
    
    return df

##  ---- 

@st.cache_data(show_spinner=False)
def load_investigacion_fmi(start_date_str, end_date_str):
    """Extractor FMI - Blogs de Investigación (Vía Coveo API con llave exacta, SOLO TÍTULO)"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # LA LLAVE MAESTRA DESCUBIERTA: IMF Blog Page
    payload = {
        "aq": "@imftype==\"IMF Blog Page\" AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        if res.status_code == 200:
            data = res.json()
            for item in data.get("results", []):
                # Extraemos el título y lo limpiamos de comillas o espacios sobrantes
                titulo_limpio = item.get("title", "").strip().strip('"').strip("'").strip()
                link = item.get("clickUri", "")
                
                # Extraemos la fecha matemática de Coveo
                raw_data = item.get("raw", {})
                raw_date = raw_data.get("date")

                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass
                
                # Si falta el título, el enlace o la fecha, descartamos el elemento
                if not titulo_limpio or not link or not parsed_date:
                    continue

                # Filtrar por fechas
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date, 
                            "Title": titulo_limpio, # <--- Se envía el título crudo sin autor
                            "Link": link, 
                            "Organismo": "FMI"
                        })
    except Exception as e:
        pass # Silencioso para no romper la aplicación

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_investigacion_bm(start_date_str, end_date_str):
    """Extractor para Investigación del BM (Filtra y excluye los que son 'Reports')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # ID exacto de la comunidad de Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id,
                'sort': 'dc.date.issued,DESC',
                'page': page,
                'size': 20
            }
            res = requests.get(base_url, headers=headers,
                               params=params, timeout=15)
            data = res.json()

            objects = data.get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects:
                break

            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})

                # Extraer Título y Fecha
                title = meta.get('dc.title', [{'value': ''}])[
                    0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[
                    0].get('value', '')

                parsed_date = None
                if date_s:
                    try:
                        parsed_date = parser.parse(date_s)
                    except:
                        pass

                if not parsed_date or parsed_date < start_date:
                    continue

                # --- NUEVO FILTRO ANTI-REPORTES ---
                # Buscamos en el abstract o en la descripción general
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])

                description = ""
                if abstract_list:
                    description = abstract_list[0].get('value', '').lower()
                elif desc_list:
                    description = desc_list[0].get('value', '').lower()

                # Si la palabra exacta "report" está en la descripción, lo saltamos
                # Usamos \b para que sea la palabra exacta y no algo como "reporting"
                if re.search(r'\breport\b', description):
                    continue
                # ----------------------------------

                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[
                    0].get('value', '')
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title,
                                "Link": link, "Organismo": "BM"})
                    items_found += 1

            if items_found == 0:
                break
            page += 1
            if page > 3:
                break  # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

## OCDE - INVESTIGACION

@st.cache_data(show_spinner=False)
def load_investigacion_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Working Papers (API oficial con paginación)"""
    import requests
    import datetime
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Investigación: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad (500 documentos máximo)
    documentos_procesados = 0
    
    print("📡 Solicitando Working Papers a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/working-papers"
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = []
            if "results" in data:
                results = data["results"]
            else:
                print(f"   ⚠️ Estructura inesperada en página {page + 1}")
                break
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            # Contar cuántos documentos del mes encontramos en esta página
            documentos_en_pagina = 0
            fecha_mas_reciente = None
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha del campo publicationDateTime
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                # Actualizar fechas extremas
                if fecha_mas_reciente is None or parsed_date > fecha_mas_reciente:
                    fecha_mas_reciente = parsed_date
                if fecha_mas_antigua is None or parsed_date < fecha_mas_antigua:
                    fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, podemos parar porque
                # los resultados están ordenados por fecha descendente
                if parsed_date < start_date:
                    # Ya no hay más documentos del mes en esta página ni en las siguientes
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    # Salimos del while principal
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite, paramos
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            # Pequeña pausa para no sobrecargar la API
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE encontrados: {documentos_procesados}")
        
    except Exception as e:
        print(f"❌ Error en load_investigacion_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Investigación - Total final: {len(df)}")
    return df

# --- SECCIÓN: DISCURSOS ---
## -- Banco de Inglaterra -- Bank of England (BoE)
@st.cache_data(show_spinner=False)
def load_discursos_boe(start_date_str, end_date_str):
    """Extractor Automático BoE - Vía RSS con formato consistente 'Autor: Título'"""
    import requests
    from bs4 import BeautifulSoup
    import pandas as pd
    import datetime
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    url = "https://www.bankofengland.co.uk/rss/speeches"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    def extract_author_from_title(title):
        """Extrae el nombre del autor del título en varios formatos"""
        autor = ""
        titulo_limpio = title
        
        # Patrón 1: "Título − speech by Autor" (con guión largo o corto)
        match = re.search(r'(?i)\s*[\-–—]\s*speech\s+by\s+(.+?)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            # Eliminar TODO desde el guión hasta el final
            titulo_limpio = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        # Patrón 2: "Speech by Autor: Título" o "Speech by Autor - Título"
        match = re.search(r'(?i)^speech\s+by\s+([^:—-]+)[:—-]\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            titulo_limpio = match.group(2).strip()
            return autor, titulo_limpio
        
        # Patrón 3: "Autor: Título" (ya está bien formateado)
        match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*:\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = match.group(2)
            return autor, titulo_limpio
        
        # Patrón 4: "Título by Autor" (sin "speech")
        match = re.search(r'(?i)\s+by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = re.sub(r'(?i)\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        return None, title

    try:
        res = requests.get(url, headers=headers, timeout=10)
        if res.status_code == 200:
            soup = BeautifulSoup(res.content, "xml")
            items = soup.find_all("item")

            for item in items:
                titulo_raw = item.find("title").text if item.find("title") else ""
                link = item.find("link").text if item.find("link") else ""
                fecha_raw = item.find("pubDate").text if item.find("pubDate") else ""

                if not titulo_raw or not link or not fecha_raw:
                    continue

                try:
                    parsed_date = parser.parse(fecha_raw)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    continue

                if start_date <= parsed_date <= end_date:
                    # Extraer autor y título limpio
                    autor, titulo_limpio = extract_author_from_title(titulo_raw)
                    
                    # LIMPIEZA DIRECTA: eliminar específicamente " − speech" o " −" al final
                    # Primero, eliminar " − speech" (con el guión especial)
                    titulo_limpio = titulo_limpio.replace(' − speech', '').replace(' - speech', '').replace('— speech', '')
                    # Luego, eliminar " −" solitario al final
                    titulo_limpio = titulo_limpio.replace(' −', '').replace(' -', '').replace('—', '')
                    # Eliminar espacios sobrantes al final
                    titulo_limpio = titulo_limpio.rstrip()
                    
                    # Construir título final en formato "Autor: Título"
                    if autor:
                        titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo_limpio
                    
                    # Limpieza final de espacios múltiples
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    titulo_final = titulo_final.strip('"').strip("'").strip()
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo_final,
                            "Link": link,
                            "Organismo": "BoE (Inglaterra)"
                        })
    except Exception as e:
        print(f"Error en load_discursos_boe: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values(by="Date", ascending=False)
    return df

## Discursos FMI - Actualización - 
@st.cache_data(show_spinner=False)
def load_discursos_fmi(start_date_str, end_date_str):
    """
    Extractor FMI - Discursos y Transcripts (Coveo API)
    """
    import datetime
    import requests
    import pandas as pd
    import re
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Discursos y Transcripts: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    # Incluir tanto SPEECHES como TRANSCRIPTS
    payload = {
        "aq": "@imftype==(\"Speech\",\"Transcript\") AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    def limpiar_titulo(titulo, speaker_name):
        """Limpia el título: elimina comillas, sufijos redundantes y el nombre del autor si está repetido"""
        if not titulo:
            return titulo
        
        # 1. Eliminar comillas
        titulo = titulo.strip('"').strip("'")
        titulo = titulo.replace('\\"', '')
        titulo = titulo.replace('"', '')
        titulo = titulo.replace("'", "")
        
        # 2. Eliminar sufijos comunes que son redundantes
        sufijos_redundantes = [
            r'\s*[-–—]\s*(?:Keynote\s+)?Speech\s+by\s+.*$',
            r'\s*[-–—]\s*(?:Opening\s+)?Remarks\s+by\s+.*$',
            r'\s*[-–—]\s*(?:Press\s+)?Briefing\s+Transcript.*$',
            r'\s*[-–—]\s*Statement\s+by\s+.*$',
            r'\s*[-–—]\s*Address\s+by\s+.*$',
            r'\s*[-–—]\s*Transcript:.*$',
        ]
        
        for patron in sufijos_redundantes:
            titulo = re.sub(patron, '', titulo, flags=re.IGNORECASE)
        
        # 3. Si el título comienza con el nombre del autor (sin cargo), eliminarlo
        if speaker_name and titulo.lower().startswith(speaker_name.lower()):
            titulo = titulo[len(speaker_name):].lstrip(': ').strip()
        
        # 4. Limpiar espacios múltiples y caracteres sobrantes
        titulo = re.sub(r'\s+', ' ', titulo).strip()
        
        # 5. Eliminar puntuación redundante al inicio
        titulo = re.sub(r'^[,:;.\s]+', '', titulo)
        
        return titulo

    try:
        print("   📡 Solicitando discursos y transcripts del FMI a Coveo API...")
        response = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)

        if response.status_code == 200:
            data = response.json()
            print(f"   ✅ Total resultados en API: {data.get('totalCount', 0)}")

            for item in data.get("results", []):
                titulo_raw = item.get("title", "").strip()
                link = item.get("clickUri", "")
                raw_date = item.get("raw", {}).get("date")
                speaker = item.get("raw", {}).get("imfspeaker", "")
                content_type = item.get("raw", {}).get("imftype", "")

                if isinstance(speaker, list) and len(speaker) > 0:
                    speaker = speaker[0]
                elif not speaker:
                    speaker = "IMF Staff"

                if not titulo_raw or not link or not raw_date:
                    continue

                try:
                    parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                except:
                    continue

                if start_date <= parsed_date <= end_date:
                    # Limpiar título usando el nombre del autor
                    titulo = limpiar_titulo(titulo_raw, speaker)
                    
                    # Construir título final
                    titulo_final = f"{speaker}: {titulo}"
                    
                    # Limpieza final
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "FMI"
                    })
                    print(f"      ✅ [{content_type}] {parsed_date.strftime('%d/%m/%Y')}: {speaker} - {titulo[:50]}...")

        else:
            print(f"   ❌ Error en API: {response.status_code}")

    except Exception as e:
        print(f"   ❌ Error: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Title'], keep='first')
        df = df.drop_duplicates(subset=['Link'], keep='first')

    print(f"📊 FMI Discursos - Total: {len(df)}")
    return df

@st.cache_data(show_spinner=False)
def load_data_ecb(start_date_str, end_date_str):
    """Extractor ECB (Europa) - FOEDB Bypass + Autor por Renglón"""
    import undetected_chromedriver as uc
    import pandas as pd
    import datetime
    import time
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://www.ecb.europa.eu/press/key/html/index.en.html"
    
    options = uc.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.binary_location = '/usr/bin/chromium'

    try:
        driver = uc.Chrome(options=options)
        driver.get(url)
        time.sleep(4)
        
        # Aceptar cookies para desbloquear inyección de datos
        js_cookie = """
        let btn = document.querySelector('.ecb-cookieConsent button.check, #cookieConsent button.check');
        if (btn) { btn.click(); return true;}
        return false;
        """
        driver.execute_script(js_cookie)
        time.sleep(5) 
        
        # Scroll para cargar el historial
        for i in range(1, 6):
            driver.execute_script(f"window.scrollTo(0, {i * 1200});")
            time.sleep(1.5)

        # Extractor de bloques visuales
        js_extract = """
        let data = [];
        let processedLinks = new Set();
        
        document.querySelectorAll('a[href*="/press/key/"]').forEach(a => {
            let title = a.innerText.trim();
            let href = a.href;
            
            if (title.length > 10 && !href.endsWith('index.en.html') && !processedLinks.has(href)) {
                processedLinks.add(href);
                
                let contextText = "";
                let dd = a.closest('dd');
                if (dd) {
                    contextText = dd.innerText;
                    if (dd.previousElementSibling && dd.previousElementSibling.tagName === 'DT') {
                        contextText += "\\n" + dd.previousElementSibling.innerText;
                    }
                } else {
                    let parent = a.closest('div[class*="result"], li, article') || a.parentElement;
                    contextText = parent ? parent.innerText : '';
                }
                
                data.push({ t: title, l: href, c: contextText });
            }
        });
        return data;
        """
        extracted = driver.execute_script(js_extract)
        
        for item in extracted:
            titulo = item['t'].strip()
            link = item['l']
            contexto = item['c'] 
            
            parsed_date = None
            date_match = re.search(r'(\d{1,2}\s+[A-Za-z]+\s+\d{4})', contexto.replace('\n', ' '))
            if date_match:
                try: parsed_date = parser.parse(date_match.group(0))
                except: pass
                
            if not parsed_date:
                m = re.search(r'/(20\d{2})/', link)
                if m: parsed_date = datetime.datetime(int(m.group(1)), 1, 1)
            
            # Búsqueda de Autor por renglón
            autor = ""
            lineas = [linea.strip() for linea in contexto.split('\n') if linea.strip()]
            
            for i, linea in enumerate(lineas):
                if titulo in linea or linea in titulo:
                    if i + 1 < len(lineas):
                        cand = lineas[i + 1]
                        if not any(x in cand.lower() for x in ['details', 'annexes', 'speech', 'interview', 'related', 'pdf']):
                            if not re.search(r'\d{4}', cand):
                                autor = clean_author_name(cand)
                    break
            
            # Fallback clásico
            if not autor:
                match_autor = re.search(r'\b(?:by|with)\s+([A-ZÀ-ÿ][A-Za-zÀ-ÿ\.\-\s]{2,35}?)(?:,|\s+at\s+|$)', titulo, re.IGNORECASE)
                if match_autor: 
                    autor = clean_author_name(match_autor.group(1))

            # Ensamblaje final del título
            titulo_limpio = titulo
            if autor:
                titulo_limpio = re.sub(rf'(?i)\s*(?:by|with)\s+{re.escape(autor)}', '', titulo_limpio)
                titulo_limpio = re.sub(r'^\s*,\s*', '', titulo_limpio).strip()
                final_title = f"{autor}: {titulo_limpio}"
            else:
                final_title = titulo
            
            if parsed_date and start_date <= parsed_date <= end_date:
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": final_title,
                        "Link": link,
                        "Organismo": "ECB (Europa)"
                    })

    except Exception as e:
        print(f"Error ECB: {e}")
    finally:
        if 'driver' in locals(): driver.quit()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Link'])
    return df

@st.cache_data(show_spinner=False)
def load_data_bis():
    urls = ["https://www.bis.org/api/document_lists/cbspeeches.json",
            "https://www.bis.org/api/document_lists/bcbs_speeches.json", "https://www.bis.org/api/document_lists/mgmtspeeches.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            data = response.json()
            for path, speech in data.get("list", {}).items():
                title = html.unescape(speech.get("short_title", ""))
                date_str = speech.get("publication_start_date", "")
                link = "https://www.bis.org" + path + \
                    (".htm" if not path.endswith(".htm") else "")
                rows.append({"Date": date_str, "Title": title,
                            "Link": link, "Organismo": "BPI"})
        except:
            continue
    df = pd.DataFrame(rows).drop_duplicates(
        subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows, page = [], 0
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str,
                  'dateTo': end_date_str, 'pageNumString': str(page)}
        try:
            response = requests.get(
                base_url, headers=headers, params=params, timeout=10)
        except:
            break
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        if not items:
            break
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            author_tag = item.find('span', class_='metadata__authors')
            author_str = clean_author_name(
                author_tag.text) if author_tag else ""
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            if data_div and data_div.find('a'):
                a_tag = data_div.find('a')
                link = "https://www.bundesbank.de" + \
                    a_tag.get('href', '') if a_tag.get(
                        'href', '').startswith('/') else a_tag.get('href', '')
                if a_tag.find('span', class_='link__label'):
                    titulo = a_tag.find(
                        'span', class_='link__label').text.strip()
            if author_str and author_str not in titulo:
                titulo = f"{author_str}: {titulo}"
            if fecha_str and titulo:
                rows.append({"Date": fecha_str, "Title": titulo,
                            "Link": link, "Organismo": "BBk (Alemania)"})
        if len(items) < 10:
            break
        page += 1
        time.sleep(0.3)
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(
            df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
    return df


## Discursos - Banco de China - PBoC
@st.cache_data(show_spinner=False)
def load_data_pboc(start_date_str, end_date_str):
    """
    Extractor PBoC (China) - Versión final con limpieza de título y cargos
    """
    import datetime
    import re
    import requests
    from bs4 import BeautifulSoup
    import pandas as pd
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 PBoC (China): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9',
    }
    
    year = start_date.year
    month = start_date.month
    
    print(f"   📡 Buscando discursos de {year}-{month:02d}...")
    
    # Intentar ambas URLs posibles
    urls_to_try = [
        "https://www.pbc.gov.cn/en/3688110/3688175/index.html",
        "https://www.pbc.gov.cn/en/3688110/3688175/index.html?page=1"
    ]
    
    for url in urls_to_try:
        try:
            response = requests.get(url, headers=headers, timeout=15, verify=False)
            
            if response.status_code != 200:
                continue
                
            # Corregir encoding
            response.encoding = 'utf-8'
            soup = BeautifulSoup(response.text, 'html.parser')
            items = soup.find_all('div', class_='prhhd1')
            
            if not items:
                # Fallback: buscar otros selectores
                items = soup.find_all('div', class_='ListR')
                if not items:
                    items = soup.find_all('li', class_='clearfix')
            
            for item in items:
                # Fecha
                date_span = item.find('span', class_='prhhdata')
                if not date_span:
                    date_span = item.find('span', class_='date')
                if not date_span:
                    continue
                
                fecha_texto = date_span.get_text(strip=True)
                
                try:
                    parsed_date = datetime.datetime.strptime(fecha_texto, '%Y-%m-%d')
                except:
                    continue
                
                if parsed_date.year != year or parsed_date.month != month:
                    continue
                
                # Enlace
                link_tag = item.find('a', href=True)
                if not link_tag:
                    continue
                
                # Extraer título
                listr_div = item.find('div', class_='ListR')
                if not listr_div:
                    listr_div = item.find('div', class_='listR')
                if not listr_div:
                    # Si no hay div específico, usar el texto del enlace
                    titulo_completo = link_tag.get_text(strip=True)
                else:
                    titulo_completo = listr_div.get_text(strip=True)
                
                # Eliminar la fecha del título
                titulo_completo = titulo_completo.replace(fecha_texto, '').strip()
                
                # ========== LIMPIEZA DEL TÍTULO ==========
                # 1. Eliminar todo después de "--Keynote" o "Keynote Speech by"
                titulo_limpio = re.split(r'--Keynote|Keynote Speech by', titulo_completo)[0].strip()
                
                # 2. Extraer autor (sin cargo)
                autor = ""
                
                # Patrón 1: "Governor Pan Gongsheng" o "Deputy Governor X"
                autor_match = re.search(r'(?:Governor|Deputy Governor|Administrator|Director|President)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', titulo_completo)
                if autor_match:
                    autor = autor_match.group(1).strip()
                
                # Patrón 2: "by Pan Gongsheng"
                if not autor:
                    name_match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', titulo_completo)
                    if name_match:
                        autor = name_match.group(1).strip()
                
                # Patrón 3: Si el título comienza con un nombre (ej. "Pan Gongsheng:")
                if not autor:
                    name_start_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s*[:：]', titulo_limpio)
                    if name_start_match:
                        autor = name_start_match.group(1).strip()
                        # Eliminar el nombre del título
                        titulo_limpio = re.sub(rf'^{re.escape(autor)}[:：]\s*', '', titulo_limpio)
                
                # 4. Construir título final
                if autor:
                    # Limpiar espacios y caracteres extraños
                    autor = re.sub(r'\s+', ' ', autor).strip()
                    titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_limpio
                
                # Limpiar apóstrofes mal codificados
                titulo_final = titulo_final.replace('â', "'").replace('â€™', "'")
                # Eliminar espacios múltiples
                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                
                link = link_tag.get('href')
                if link and link.startswith('/'):
                    link = f"https://www.pbc.gov.cn{link}"
                elif not link:
                    continue
                
                # Verificar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "PBoC (China)"
                    })
                    print(f"      ✅ {parsed_date.strftime('%d/%m/%Y')}: {titulo_final[:60]}...")
                    
        except Exception as e:
            print(f"   ⚠️ Error con URL {url}: {e}")
            continue
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Title'], keep='first')
        df = df.drop_duplicates(subset=['Link'], keep='first')
    
    print(f"📊 PBoC (China) - Total: {len(df)} discursos")
    return df


@st.cache_data(show_spinner=False)
def load_data_fed(anios_num):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for year in anios_num:
        url = f"https://www.federalreserve.gov/newsevents/{year}-speeches.htm"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            if res.status_code == 404:
                url = "https://www.federalreserve.gov/newsevents/speeches.htm"
                res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a_tag in soup.find_all('a', href=True):
                if '/newsevents/speech/' in a_tag['href']:
                    link = "https://www.federalreserve.gov" + \
                        a_tag['href'] if a_tag['href'].startswith(
                            '/') else a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    parent = a_tag.find_parent(
                        'div', class_='row') or a_tag.parent
                    text = parent.get_text(separator=' | ', strip=True)
                    date_m = re.search(
                        r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s\d{1,2},\s\d{4})', text)
                    if date_m:
                        try:
                            parsed_date = parser.parse(date_m.group(1))
                            if parsed_date.year not in anios_num:
                                continue
                            rows.append({"Date": parsed_date, "Title": titulo,
                                        "Link": link, "Organismo": "Fed (Estados Unidos)"})
                        except:
                            pass
        except:
            pass
    df = pd.DataFrame(rows).drop_duplicates(
        subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


## Banco de Francia - BDF - Discursos 
@st.cache_data(show_spinner=False)
def load_data_bdf(start_date_str, end_date_str):
    """Extractor Banco de Francia (BdF) - Discursos del Gobernador (Versión Selenium)"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BdF (Francia) - Selenium: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    
    # URL principal con el filtro de discursos del Gobernador
    url = "https://www.banque-france.fr/en/governor-interventions?category%5B7052%5D=7052"
    
    # Configuración de Selenium
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)
    
    try:
        print(f"📡 Iniciando Selenium para BdF...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        print(f"   Navegando a: {url}")
        driver.get(url)
        
        # Esperar a que cargue el contenido principal
        time.sleep(5)
        
        # Scroll para activar lazy loading si existe
        driver.execute_script("window.scrollTo(0, 1000);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 2000);")
        time.sleep(2)
        
        # Extraer el HTML ya renderizado
        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')
        
        # Buscar los cards de discursos
        cards = soup.find_all('div', class_=lambda c: c and 'card' in c if c else False)
        
        # Si no encuentra cards, buscar directamente con selectores más específicos
        if not cards:
            cards = soup.find_all('div', class_='card')
        
        print(f"   📚 Cards encontrados: {len(cards)}")
        
        # Si aún no hay cards, buscar artículos
        if not cards:
            cards = soup.find_all('article')
            print(f"   📚 Artículos encontrados: {len(cards)}")
        
        # Mapeo de meses en inglés para fechas como "2nd of April 2026"
        meses_map = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }
        
        items_found = 0
        for card in cards:
            try:
                # === 1. EXTRAER TÍTULO Y ENLACE ===
                title_elem = None
                link = None
                
                # Buscar h3 con clase card__title o similar
                title_h3 = card.find('h3', class_=lambda c: c and 'card__title' in c if c else False)
                if not title_h3:
                    title_h3 = card.find('h3')
                
                if title_h3:
                    a_tag = title_h3.find('a')
                    if a_tag:
                        title_elem = a_tag
                        link = a_tag.get('href', '')
                
                if not title_elem:
                    # Buscar cualquier enlace con texto largo
                    for a in card.find_all('a', href=True):
                        texto = a.get_text(strip=True)
                        if len(texto) > 20:
                            title_elem = a
                            link = a.get('href', '')
                            break
                
                if not title_elem or not link:
                    continue
                
                titulo = title_elem.get_text(strip=True)
                
                # Limpiar título (eliminar saltos de línea y espacios extra)
                titulo = re.sub(r'\s+', ' ', titulo).strip()
                
                # === NUEVO: Eliminar comillas tipográficas del título original ===
                # Eliminar comillas dobles inglesas y españolas (apertura y cierre)
                titulo = titulo.replace('“', '').replace('”', '').replace('"', '').replace('«', '').replace('»', '')
                # Eliminar comillas simples si existen
                titulo = titulo.replace("'", "")

                # Construir URL absoluta
                if link.startswith('/'):
                    link = "https://www.banque-france.fr" + link
                
                # === 2. EXTRAER FECHA ===
                date_elem = None
                date_text = None
                
                # Buscar div con clase card__date
                date_div = card.find('div', class_=lambda c: c and 'card__date' in c if c else False)
                if date_div:
                    date_text = date_div.get_text(strip=True)
                else:
                    # Buscar cualquier elemento con clase que contenga 'date'
                    date_elem = card.find(class_=re.compile(r'date', re.I))
                    if date_elem:
                        date_text = date_elem.get_text(strip=True)
                
                if not date_text:
                    # Buscar en el texto del card
                    card_text = card.get_text()
                    date_match = re.search(r'(\d{1,2}(?:st|nd|rd|th)?\s+of\s+[A-Za-z]+\s+\d{4})', card_text, re.IGNORECASE)
                    if date_match:
                        date_text = date_match.group(1)
                
                if not date_text:
                    continue
                
                # Limpiar fecha: eliminar "st", "nd", "rd", "th" y "of"
                date_text = re.sub(r'(\d+)(st|nd|rd|th)\s+of\s+', r'\1 ', date_text, flags=re.IGNORECASE)
                date_text = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', date_text)
                date_text = date_text.strip()
                
                # Parsear fecha
                parsed_date = None
                try:
                    # Intentar parsear formatos como "2 April 2026" o "April 2, 2026"
                    parsed_date = parser.parse(date_text)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    # Fallback: extraer manualmente
                    match = re.search(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', date_text, re.IGNORECASE)
                    if not match:
                        match = re.search(r'([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})', date_text, re.IGNORECASE)
                    
                    if match:
                        groups = match.groups()
                        if len(groups) == 3:
                            # Determinar si el primer grupo es día o mes
                            if groups[0].isdigit():
                                dia = int(groups[0])
                                mes_str = groups[1].lower()
                                año = int(groups[2])
                            else:
                                mes_str = groups[0].lower()
                                dia = int(groups[1])
                                año = int(groups[2])
                            
                            mes_num = meses_map.get(mes_str, 1)
                            try:
                                parsed_date = datetime.datetime(año, mes_num, min(dia, 28))
                            except:
                                parsed_date = datetime.datetime(año, mes_num, 1)
                
                if not parsed_date:
                    continue
                
                # === 3. FILTRAR POR FECHA ===
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # === 4. VERIFICAR DUPLICADOS ===
                if not any(r['Link'] == link for r in rows):
                    # === NUEVO: Extraer autor desde la página del discurso ===
                    autor = None
                    # Solo intentar si el título no tiene ya formato "Nombre:"
                    if not re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+:', titulo):
                        try:
                            headers_page = {'User-Agent': 'Mozilla/5.0'}
                            page_response = requests.get(link, headers=headers_page, timeout=10)
                            if page_response.status_code == 200:
                                soup_page = BeautifulSoup(page_response.text, 'html.parser')
                                page_text = soup_page.get_text()
                                
                                # === CÓDIGO CORREGIDO ===
                                # Incluir letras acentuadas y cedilla: A-Za-zÀ-ÿç
                                match = re.search(r'Speech by ([A-Za-zÀ-ÿç\s]+?)(?:\s+Governor|\s+of|\s*$)', page_text)
                                if not match:
                                    # Fallback: capturar primeras palabras después de "Speech by"
                                    match = re.search(r'Speech by ([A-ZÀ-ÿ][a-zÀ-ÿç]+(?:\s+[A-Za-zÀ-ÿç]+)?(?:\s+[a-zÀ-ÿç]+)?(?:\s+[A-Za-zÀ-ÿç]+)?)', page_text)
                                
                                if match:
                                    autor = match.group(1).strip()
                                    # Limpiar espacios extra
                                    autor = re.sub(r'\s+', ' ', autor)
                                    print(f"      📝 Autor encontrado: {autor}")
                        except:
                            pass
                    
                    if autor:
                        # Limpiar título: eliminar comillas y espacios extra
                        titulo_limpio = titulo.strip()
                        # Eliminar comillas dobles inglesas y españolas (apertura y cierre)
                        for char in ['"', "'", '“', '”', '«', '»']:
                            if titulo_limpio.startswith(char) and titulo_limpio.endswith(char):
                                titulo_limpio = titulo_limpio[1:-1]
                                break
                        
                        # Verificar si el autor ya está al inicio del título (evitar duplicados)
                        if titulo_limpio.lower().startswith(autor.lower()):
                            titulo_final = titulo_limpio  # No añadir autor duplicado
                        else:
                            titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "BdF (Francia)"
                    })
                    items_found += 1
                    print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo_final[:60]}...")
                    items_found += 1
                
            except Exception as e:
                print(f"   ⚠️ Error procesando card: {e}")
                continue
        
        print(f"   📊 Documentos encontrados en BdF: {items_found}")
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error en load_data_bdf: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 BdF (Francia) - Total final: {len(df)}")
    return df


## Banco de Canadá (BOC , boc)
@st.cache_data(show_spinner=False)
def load_data_boc(start_date_str, end_date_str):
    """Extractor Banco de Canadá (BoC) - Versión corregida con prioridad a conferencias"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser
    import requests

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BoC (Canadá): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    page = 1
    
    def limpiar_titulo(titulo):
        """Limpia el título de texto basura"""
        titulo = re.sub(r'(?i)^Speech\s*[:\-]\s*', '', titulo)
        titulo = re.sub(r'(?i)^Remarks\s*[:\-]\s*', '', titulo)
        titulo = re.sub(r'(?i)^Opening\s+statement\s*[:\-]\s*', '', titulo)
        titulo = re.sub(r'(?i)^Fireside\s+chat\s*[:\-]\s*', '', titulo)
        titulo = re.sub(r'(?i)^Press\s+Conference\s*[:\-]\s*', '', titulo)
        return titulo.strip()
    
    def extraer_autor_desde_html(soup_page, titulo_raw, url):
        """Extrae el autor basado en el contenido específico de cada página"""
        
        page_text = soup_page.get_text()
        
        # === 1. PRIORIDAD MÁXIMA: Conferencias de prensa (webcasts) ===
        # Verificar si es un webcast (video) por la URL o el título
        is_webcast = 'multimedia' in url or 'webcast' in titulo_raw.lower()
        
        if 'press conference' in titulo_raw.lower():
            if is_webcast:
                # Solo para webcasts/videos, asignar ambos autores
                if 'Tiff Macklem' in page_text and 'Carolyn Rogers' in page_text:
                    return "Tiff Macklem and Carolyn Rogers"
                if 'Tiff Macklem' in page_text:
                    return "Tiff Macklem"
                if 'Carolyn Rogers' in page_text:
                    return "Carolyn Rogers"
                return "Bank of Canada"
            else:
                # Para textos (opening statements), solo Tiff Macklem
                if 'Tiff Macklem' in page_text:
                    return "Tiff Macklem"
                if 'Carolyn Rogers' in page_text:
                    return "Carolyn Rogers"
                return "Bank of Canada"
        
        # === 2. Opening statements (no son conferencias de prensa) ===
        if 'opening statement' in titulo_raw.lower() or 'opening' in titulo_raw.lower():
            if 'Tiff Macklem' in page_text:
                return "Tiff Macklem"
            if 'Carolyn Rogers' in page_text:
                return "Carolyn Rogers"
        
        # === 3. Sharon Kozicki ===
        if 'Sharon Kozicki' in page_text and ('Deputy Governor' in page_text or 'speech summary' in page_text.lower()):
            return "Sharon Kozicki"
        
        # === 4. Tiff Macklem ===
        if 'Tiff Macklem' in page_text and ('Governor' in page_text or 'remarks' in page_text.lower()):
            return "Tiff Macklem"
        
        # === 5. Carolyn Rogers ===
        if 'Carolyn Rogers' in page_text and ('Senior Deputy Governor' in page_text or 'speech summary' in page_text.lower()):
            return "Carolyn Rogers"
        
        # === 6. Fallback a media-authors ===
        author_span = soup_page.find('span', class_='media-authors')
        if author_span:
            author_link = author_span.find('a')
            if author_link:
                return author_link.text.strip()
        
        return None
    
    while True:
        try:
            if page == 1:
                url = "https://www.bankofcanada.ca/press/speeches/"
            else:
                url = f"https://www.bankofcanada.ca/press/speeches/page/{page}/"
            
            print(f"📄 Procesando página {page}: {url}")
            
            chrome_options = Options()
            chrome_options.add_argument('--headless=new')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1920,1080')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
            
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            time.sleep(5)
            
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            driver.quit()
            
            articles = soup.find_all('div', class_=lambda c: c and ('mtt-result' in c or 'media' in c or 'entry' in c))
            
            if not articles:
                articles = soup.find_all('article')
            
            if not articles:
                print(f"   📭 No hay más artículos en página {page}")
                break
            
            print(f"   📚 Artículos encontrados: {len(articles)}")
            items_found = 0
            
            for art in articles:
                h3 = art.find('h3', class_='media-heading')
                if not h3:
                    h3 = art.find('h3')
                if not h3:
                    continue
                
                a_tag = h3.find('a')
                if not a_tag:
                    continue
                
                titulo_raw = a_tag.text.strip()
                link = a_tag['href']
                
                date_elem = art.find('span', class_='media-date')
                if not date_elem:
                    date_elem = art.find('time')
                if not date_elem:
                    date_elem = art.find(class_=re.compile(r'date', re.I))
                
                if not date_elem:
                    continue
                
                try:
                    fecha_texto = date_elem.text.strip()
                    parsed_date = parser.parse(fecha_texto)
                except:
                    continue
                
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                if any(r['Link'] == link for r in rows):
                    continue
                
                print(f"   🔍 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo_raw[:50]}...")
                
                autor = None
                
                try:
                    headers = {'User-Agent': 'Mozilla/5.0'}
                    response = requests.get(link, headers=headers, timeout=10)
                    if response.status_code == 200:
                        soup_page = BeautifulSoup(response.text, 'html.parser')
                        autor = extraer_autor_desde_html(soup_page, titulo_raw, link)
                        if autor:
                            print(f"      📝 Autor encontrado: {autor}")
                except Exception as e:
                    print(f"      ⚠️ Error obteniendo página: {e}")
                
                titulo_limpio = limpiar_titulo(titulo_raw)
                
                if autor:
                    # Limpiar "Governor" del autor si está presente
                    autor_limpio = re.sub(r'^Governor\s+', '', autor)
                    titulo_final = f"{autor_limpio}: {titulo_limpio}"
                else:
                    titulo_final = titulo_limpio
                
                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                
                rows.append({
                    "Date": parsed_date,
                    "Title": titulo_final,
                    "Link": link,
                    "Organismo": "BoC (Canadá)"
                })
                items_found += 1
                print(f"      ✅ Agregado: {titulo_final[:80]}...")
            
            if items_found == 0:
                break
            
            page += 1
            time.sleep(1)
            
        except Exception as e:
            print(f"   ⚠️ Error en página {page}: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 BoC (Canadá) - Total final: {len(df)}")
    return df


## Conversor de Nombre (Nombre, Apellido) para Autores del Banco de Japón  
def convertir_nombre_japones(nombre):
    """
    Convierte nombre japonés (apellido primero) a formato occidental.
    
    Ejemplos:
    - "UEDA Kazuo" -> "Kazuo UEDA"
    - "UEDA Kazuo San" -> "Kazuo San UEDA"
    - "KURODA Haruhiko" -> "Haruhiko KURODA"
    - "AMAMIYA Masayoshi" -> "Masayoshi AMAMIYA"
    """
    if not nombre:
        return nombre
    
    partes = nombre.split()
    if len(partes) < 2:
        return nombre
    
    # La primera palabra es el apellido, el resto es el nombre
    apellido = partes[0]
    nombre_pila = " ".join(partes[1:])
    
    # Formato occidental: "Nombre Apellido"
    return f"{nombre_pila} {apellido}"

## Bank of Japan (BOJ - boj) - Discursos
@st.cache_data(show_spinner=False)
def load_data_boj(start_date_str, end_date_str):
    base_url = "https://www.boj.or.jp/en/about/press/index.htm"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BoJ (Japón): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    try:
        response = requests.get(base_url, headers=headers, timeout=12)
        soup = BeautifulSoup(response.text, 'html.parser')
        table = soup.find('table', class_='js-tbl')
        if table:
            for tr in table.find('tbody').find_all('tr'):
                tds = tr.find_all('td')
                if len(tds) < 3:
                    continue
                
                # === 1. EXTRAER FECHA ===
                try:
                    fecha_texto = tds[0].get_text(strip=True).replace('\xa0', ' ')
                    parsed_date = parser.parse(fecha_texto)
                except:
                    continue
                
                # Filtrar por rango de fechas
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # === 2. EXTRAER AUTOR (NUEVO) ===
                autor_raw = tds[1].get_text(strip=True)
                autor = None
                if autor_raw:
                    # Limpiar el autor: eliminar "Governor", "Deputy Governor", etc.
                    # Ejemplo: "UEDA Kazuo, Governor" -> "UEDA Kazuo"
                    autor = re.sub(r',\s*(Governor|Deputy Governor|Member of the Policy Board)$', '', autor_raw)
                    # Limpiar espacios extra
                    autor = autor.strip()
                    # === CONVERTIR NOMBRE JAPONÉS A FORMATO OCCIDENTAL ===
                    autor = convertir_nombre_japones(autor)
                
                # === 3. EXTRAER TÍTULO Y ENLACE ===
                a_tag = tds[2].find('a', href=True)
                if not a_tag:
                    continue
                
                titulo_raw = a_tag.get_text(strip=True).strip('"')
                link = "https://www.boj.or.jp" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                
                # === 4. CONSTRUIR TÍTULO FINAL CON AUTOR ===
                if autor:
                    # Limpiar título: eliminar el nombre del autor si está repetido
                    titulo_limpio = titulo_raw
                    # Si el título comienza con el nombre del autor, lo removemos
                    if titulo_limpio.startswith(autor.split(',')[0]):
                        titulo_limpio = re.sub(r'^[^:：]+[:：]\s*', '', titulo_limpio)
                    
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_raw
                
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo_final, 
                    "Link": link, 
                    "Organismo": "BoJ (Japón)"
                })
                print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo_final[:60]}...")
    except Exception as e:
        print(f"⚠️ Error en load_data_boj: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    print(f"📊 BoJ (Japón) - Total final: {len(df)}")
    return df


## Discursos - CEF 
@st.cache_data(show_spinner=False)
def load_data_cef(start_date_str, end_date_str):
    """
    Extractor CEF (FSB) - SOLO Discursos y Statements
    Con manejo robusto de timeouts y fallbacks para autor
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 CEF (FSB): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }
    
    rows = []
    page = 1
    
    def es_discurso(url, titulo):
        """Determina si una página es un discurso"""
        titulo_lower = titulo.lower()
        url_lower = url.lower()
        
        # Excluir comunicados de prensa puros
        if re.match(r'^(fsb publishes|fsb warns|fsb chair warns)(?!.*(speech|keynote|summit))', titulo_lower):
            return False
        
        # Incluir por URL
        if any(keyword in url_lower for keyword in ['/speech/', '/statement/', '/remarks/']):
            return True
        
        # Incluir por palabras clave en título
        if any(keyword in titulo_lower for keyword in [
            'speech', 'keynote', 'remarks', 'statement', 'foreword', 
            'address', 'testimony', 'opening remarks', 'closing remarks'
        ]):
            return True
        
        # Incluir si menciona autoridades del FSB
        if any(title_word in titulo_lower for title_word in ['fsb chair', 'secretary general', 'deputy governor']):
            return True
        
        return False
    
    def inferir_autor_desde_titulo(titulo):
        """Infiere el autor basándose en el título cuando no se puede acceder a la página"""
        titulo_lower = titulo.lower()
        
        # Palabras clave que indican quién es el autor
        if 'fsb chair' in titulo_lower or 'chair' in titulo_lower:
            return 'Andrew Bailey'
        if 'secretary general' in titulo_lower:
            return 'John Schindler'
        if 'deputy governor' in titulo_lower:
            # Podría ser varios, pero intentamos extraer del contexto
            if 'john schindler' in titulo_lower:
                return 'John Schindler'
            return 'FSB Deputy Governor'
        
        return None
    
    def extraer_autor_y_titulo_desde_pagina(url, titulo_lista):
        """Extrae el autor y el título limpio de la página individual con manejo de timeouts"""
        autor = None
        titulo_limpio = titulo_lista
        
        try:
            # Timeout más generoso y reintento
            time.sleep(0.5)
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code != 200:
                # Fallback: inferir autor del título
                autor = inferir_autor_desde_titulo(titulo_lista)
                return autor, titulo_limpio
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # === OBTENER TÍTULO CORRECTO DEL <h1> ===
            h1_tag = soup.find('h1')
            if h1_tag:
                titulo_limpio = h1_tag.get_text(strip=True)
                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
            
            # === EXTRAER AUTOR ===
            # Método 1: Buscar en el bloque blockquote
            blockquote = soup.find('blockquote')
            if blockquote:
                texto = blockquote.get_text()
                match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s+(?:the\s+)?(?:Chair|Secretary General|Deputy Governor|Governor)', texto)
                if match:
                    autor = match.group(1).strip()
                
                if not autor:
                    match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', texto)
                    if match:
                        autor = match.group(1).strip()
            
            # Método 2: Buscar en meta tags de perfil
            if not autor:
                meta_profile = soup.find('meta', attrs={'name': 'fsb_profile_post'})
                if meta_profile:
                    profile_value = meta_profile.get('content', '').lower()
                    nombres = {
                        'andrew-bailey': 'Andrew Bailey',
                        'john-schindler': 'John Schindler',
                        'klaas-knot': 'Klaas Knot',
                        'martin-moloney': 'Martin Moloney'
                    }
                    for key, name in nombres.items():
                        if key in profile_value:
                            autor = name
                            break
            
            # Método 3: Si el título contiene "FSB Chair", el autor es Andrew Bailey
            if not autor and ('FSB Chair' in titulo_limpio or 'Chair' in titulo_limpio):
                autor = 'Andrew Bailey'
            
            # Método 4: Buscar en el contenido del artículo
            if not autor:
                article = soup.find('article')
                if article:
                    text = article.get_text()
                    match = re.search(r'Speech\s+by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text, re.IGNORECASE)
                    if match:
                        autor = match.group(1).strip()
            
            # Si aún no hay autor, intentar inferir del título
            if not autor:
                autor = inferir_autor_desde_titulo(titulo_limpio)
            
            return autor, titulo_limpio
            
        except requests.exceptions.Timeout:
            print(f"      ⚠️ Timeout al acceder a {url}, infiriendo autor del título...")
            autor = inferir_autor_desde_titulo(titulo_lista)
            return autor, titulo_limpio
        except Exception as e:
            print(f"      ⚠️ Error: {e}")
            autor = inferir_autor_desde_titulo(titulo_lista)
            return autor, titulo_limpio
    
    while True:
        try:
            if page == 1:
                url = "https://www.fsb.org/press/speeches-and-statements/"
            else:
                url = f"https://www.fsb.org/press/speeches-and-statements/page/{page}/"
            
            print(f"📄 Procesando página {page}: {url}")
            
            response = requests.get(url, headers=headers, timeout=20)
            if response.status_code != 200:
                print(f"   ❌ Error HTTP: {response.status_code}")
                break
            
            soup = BeautifulSoup(response.text, 'html.parser')
            items = soup.find_all('div', class_='post-excerpt')
            
            if not items:
                items = soup.find_all('div', class_=lambda c: c and 'post-excerpt' in c if c else False)
            
            if not items:
                print(f"   📭 No se encontraron más elementos en página {page}")
                break
            
            print(f"   📚 Elementos encontrados: {len(items)}")
            items_found = 0
            
            for item in items:
                try:
                    title_elem = item.find('h3')
                    if not title_elem:
                        title_elem = item.find('div', class_='post-title')
                    
                    if not title_elem:
                        continue
                    
                    a_tag = title_elem.find('a')
                    if not a_tag:
                        continue
                    
                    titulo_raw = a_tag.get_text(strip=True)
                    link = a_tag.get('href', '')
                    
                    if not link:
                        continue
                    
                    date_elem = item.find('div', class_='post-date')
                    if not date_elem:
                        date_elem = item.find('span', class_='post-date')
                    
                    if not date_elem:
                        continue
                    
                    fecha_texto = date_elem.get_text(strip=True)
                    
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    if any(r['Link'] == link for r in rows):
                        continue
                    
                    print(f"   🔍 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo_raw[:50]}...")
                    
                    if not es_discurso(link, titulo_raw):
                        print(f"      ⏭️ Excluido (no es discurso): {titulo_raw[:50]}...")
                        continue
                    
                    # === EXTRAER AUTOR Y TÍTULO ===
                    autor, titulo_limpio = extraer_autor_y_titulo_desde_pagina(link, titulo_raw)
                    
                    # === CONSTRUIR TÍTULO FINAL ===
                    if autor and titulo_limpio:
                        # Verificar si el autor ya está al inicio del título
                        if not titulo_limpio.lower().startswith(autor.lower()):
                            titulo_final = f"{autor}: {titulo_limpio}"
                        else:
                            titulo_final = titulo_limpio
                    else:
                        titulo_final = titulo_limpio
                    
                    # Limpieza mínima
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    titulo_final = titulo_final.replace('â', "'").replace('â€™', "'")
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "CEF"
                    })
                    items_found += 1
                    print(f"      ✅ Discurso: {titulo_final[:80]}...")
                    
                except Exception as e:
                    print(f"   ⚠️ Error procesando item: {e}")
                    continue
            
            print(f"   📊 Discursos en página {page}: {items_found}")
            
            # Si no encontramos discursos en 2 páginas consecutivas, paramos
            if items_found == 0 and page > 2:
                break
            
            page += 1
            time.sleep(1.5)  # Pausa más larga entre páginas
            
        except requests.exceptions.Timeout:
            print(f"   ⏱️ Timeout en página {page}, continuando...")
            page += 1
            time.sleep(3)
            continue
        except Exception as e:
            print(f"❌ Error en página {page}: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n📊 CEF (FSB) - Total final: {len(df)} discursos")
    return df

## - Discursos - Banco de España - 
@st.cache_data(show_spinner=False)
def load_data_bde(start_date_str, end_date_str):
    """Extractor Banco de España - Versión con extracción de nombres reales desde PDF"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from PyPDF2 import PdfReader
    import io
    import requests
    import datetime
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BdE (España): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://www.bde.es/wbe/en/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/"
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    
    def extraer_autor_y_cargo_desde_pdf(pdf_url):
        """Extrae el nombre y cargo del autor desde el PDF"""
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(pdf_url, headers=headers, timeout=15)
            if response.status_code != 200:
                return None, None
            
            pdf_file = io.BytesIO(response.content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for i in range(min(3, len(reader.pages))):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if not text:
                return None, None
            
            lineas = text.split('\n')
            nombre = None
            cargo = None
            
            for i, linea in enumerate(lineas):
                linea_limpia = linea.strip()
                
                if re.search(r'Governor|Gobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'Deputy Governor|Subgobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Deputy Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'Subgobernadora', linea_limpia, re.IGNORECASE):
                    cargo = "Subgobernadora"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'D\.G\.|Director General', linea_limpia, re.IGNORECASE):
                    cargo = "Director General"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
            
            if not nombre:
                for linea in lineas[:15]:
                    linea_limpia = linea.strip()
                    if re.match(r'^[A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){1,3}$', linea_limpia):
                        if not any(palabra in linea_limpia for palabra in ['DIRECTOR', 'GENERAL', 'DEPARTAMENTO', 'SECRETARÍA', 'MINISTERIO', 'GOBIERNO', 'BANCO', 'ESPAÑA', 'MADRID']):
                            nombre = linea_limpia
                            break
            
            if nombre:
                nombre = ' '.join(nombre.split())
                nombre = nombre.title()
                nombre = re.sub(r'\bDe\b', 'de', nombre)
                nombre = re.sub(r'\bY\b', 'y', nombre)
                return nombre, cargo
            
            return None, None
            
        except Exception as e:
            print(f"      ⚠️ Error extrayendo del PDF: {e}")
            return None, None

    try:
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(8)

        js_script = """
        let data = [];
        let results = document.querySelectorAll('.block-search-result, .block-search-result--image');
        results.forEach(el => {
            let titleEl = el.querySelector('.block-search-result__title, a');
            let dateEl = el.querySelector('.block-search-result__date');
            let linkEl = el.querySelector('a');
            if (titleEl && dateEl && linkEl) {
                data.push({
                    title: titleEl.innerText,
                    dateText: dateEl.innerText,
                    link: linkEl.href
                });
            }
        });
        return data;
        """
        extracted = driver.execute_script(js_script)
        driver.quit()

        print(f"   📚 Discursos encontrados: {len(extracted)}")

        for idx, item in enumerate(extracted):
            raw_title = item['title'].strip()
            raw_date_str = item['dateText'].strip()
            page_link = item['link']
            
            if not raw_title or not raw_date_str:
                continue

            parsed_date = None
            try:
                parsed_date = datetime.datetime.strptime(raw_date_str, '%d/%m/%Y')
            except:
                match = re.search(r'(\d{2}/\d{2}/\d{4})', raw_date_str)
                if match:
                    parsed_date = datetime.datetime.strptime(match.group(1), '%d/%m/%Y')

            if parsed_date and start_date <= parsed_date <= end_date:
                print(f"   🔍 Procesando ({idx+1}/{len(extracted)}): {parsed_date.strftime('%Y-%m-%d')}")
                
                try:
                    page_response = requests.get(page_link, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
                    if page_response.status_code == 200:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(page_response.text, 'html.parser')
                        pdf_link = None
                        for a in soup.find_all('a', href=True):
                            if a['href'].endswith('.pdf'):
                                pdf_link = a['href']
                                if pdf_link.startswith('/'):
                                    pdf_link = "https://www.bde.es" + pdf_link
                                break
                        
                        if pdf_link:
                            print(f"      📄 PDF encontrado, extrayendo autor...")
                            autor, cargo = extraer_autor_y_cargo_desde_pdf(pdf_link)
                            # Dentro de load_data_bde(), después de encontrar el autor
                            if autor:
                                titulo_limpio = raw_title
                                
                                # ========== NUEVA LIMPIEZA MEJORADA ==========
                                # Eliminar patrones comunes de cargo (en español e inglés)
                                patrones_cargo = [
                                    r'D\.G\.\s*Econom[íi]a\.\s*',      # D.G. Economía. o D.G. Economics.
                                    r'D\.G\.\s*Economics\.\s*',         # D.G. Economics.
                                    r'Deputy\s*Governor\.\s*',          # Deputy Governor.
                                    r'Governor\.\s*',                   # Governor.
                                    r'Subgobernador[a]?\.\s*',          # Subgobernadora. o Subgobernador.
                                    r'Director\s*General\.\s*',         # Director General.
                                    r'Head\s*of\s*\w+\.\s*',            # Head of Department.
                                    r'Director\.\s*',                   # Director.
                                    r'Chief\s*Economist\.\s*',          # Chief Economist.
                                    r'Gerente\s*General\.\s*',          # Gerente General.
                                    r'Vicepresident[ae]\.\s*',          # Vicepresidenta. o Vicepresidente.
                                    r'President[ae]\.\s*',              # Presidenta. o Presidente.
                                ]
                                
                                for patron in patrones_cargo:
                                    titulo_limpio = re.sub(patron, '', titulo_limpio, flags=re.IGNORECASE)
                                
                                # También eliminar cualquier texto entre paréntesis que parezca un cargo
                                titulo_limpio = re.sub(r'\s*\([^)]*(?:D\.G\.|Governor|Director|Econom[íi]a)[^)]*\)\s*', ' ', titulo_limpio, flags=re.IGNORECASE)
                                
                                # Limpiar espacios múltiples y puntos al inicio
                                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
                                titulo_limpio = re.sub(r'^\.\s*', '', titulo_limpio)
                                
                                # Construir título final
                                titulo_final = f"{autor}: {titulo_limpio}"
                                
                                # Limpieza adicional: eliminar " : " si el título está vacío
                                titulo_final = re.sub(r':\s*$', '', titulo_final)
                                
                                print(f"      📝 Título limpio: {titulo_final[:80]}...")
                            
                            else:
                                print(f"      ⚠️ No se pudo extraer autor, usando formato estándar")
                                titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    else:
                        titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                        
                except Exception as e:
                    print(f"      ⚠️ Error accediendo a la página: {e}")
                    titulo_final = re.sub(r'\.\s+', ': ', raw_title, count=1)
                
                if not any(r['Link'] == page_link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": page_link,
                        "Organismo": "BdE (España)"
                    })
                    print(f"      ✅ Agregado: {titulo_final[:80]}...")

    except Exception as e:
        print(f"❌ Error BDE: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BdE (España) - Total final: {len(df)}")
    return df

# ==========================================
# NUEVAS FUNCIONES PARA BID (bypass Cloudflare)
# ==========================================

@st.cache_data(show_spinner=False)
def load_investigacion_bid_cloudscraper(start_date_str, end_date_str):
    """
    Extrae Working Papers usando cloudscraper (bypass Cloudflare)
    """
    try:
        import cloudscraper
    except ImportError:
        print("❌ cloudscraper no instalado. Ejecuta: pip install cloudscraper")
        return pd.DataFrame()
    
    from bs4 import BeautifulSoup
    import datetime
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Cloudscraper: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # Crear scraper con configuraciones específicas
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    # URLs a probar
    urls_to_try = [
        "https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers",
        "https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo"
    ]
    
    for url in urls_to_try:
        lang = "en" if "en?" in url else "es"
        try:
            print(f"📡 Accediendo a {url[:60]}...")
            response = scraper.get(url, timeout=30)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extraer artículos
                articles = soup.find_all('div', class_='views-row')
                print(f"   📚 Artículos encontrados: {len(articles)}")
                
                for article in articles:
                    # Extraer título y link
                    title_elem = article.find('div', class_='views-field-field-title')
                    if not title_elem:
                        continue
                    
                    a_tag = title_elem.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Extraer fecha
                    date_elem = article.find('div', class_='views-field-field-date-issued-text')
                    if date_elem:
                        date_text = date_elem.get_text(strip=True)
                        # Parsear fecha "Mar 2026"
                        match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año = match.groups()
                            meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                                   'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12}
                            mes = meses.get(mes_str, 1)
                            parsed_date = datetime.datetime(int(año), mes, 1)
                            
                            if start_date <= parsed_date <= end_date:
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                                })
                                print(f"      ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
            
        except Exception as e:
            print(f"⚠️ Error en {lang}: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Cloudscraper - Total: {len(df)} documentos")
    return df


@st.cache_data(show_spinner=False)
def load_investigacion_bid_selenium_fallback(start_date_str, end_date_str):
    """
    Fallback: Extrae Working Papers con Selenium + delay largo
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Selenium Fallback: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    chrome_options = Options()
    chrome_options.add_argument('--headless=new')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-blink-features=AutomationControlled')
    chrome_options.add_argument('--window-size=1920,1080')
    
    urls = [
        ("https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers", "en"),
        ("https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo", "es")
    ]
    
    for url, lang in urls:
        driver = None
        try:
            print(f"📡 Accediendo con Selenium a {url[:60]}...")
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            
            # ⚠️ CLAVE: Esperar a que Cloudflare resuelva
            print("   ⏳ Esperando 20 segundos para Cloudflare...")
            time.sleep(20)
            
            # Scroll para cargar contenido
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(3)
            
            # Extraer usando BeautifulSoup
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            articles = soup.find_all('div', class_='views-row')
            print(f"   📚 Artículos encontrados: {len(articles)}")
            
            for article in articles:
                title_elem = article.find('div', class_='views-field-field-title')
                if not title_elem:
                    continue
                
                a_tag = title_elem.find('a')
                if not a_tag:
                    continue
                
                titulo = a_tag.get_text(strip=True)
                link = a_tag.get('href')
                if link and not link.startswith('http'):
                    link = "https://publications.iadb.org" + link
                
                date_elem = article.find('div', class_='views-field-field-date-issued-text')
                if date_elem:
                    date_text = date_elem.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                    if match:
                        mes_str, año = match.groups()
                        meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                               'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12,
                               'ene':1, 'feb':2, 'mar':3, 'abr':4, 'may':5, 'jun':6,
                               'jul':7, 'ago':8, 'sep':9, 'oct':10, 'nov':11, 'dic':12}
                        mes = meses.get(mes_str, 1)
                        parsed_date = datetime.datetime(int(año), mes, 1)
                        
                        if start_date <= parsed_date <= end_date:
                            rows.append({
                                "Date": parsed_date,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                            })
            
        except Exception as e:
            print(f"⚠️ Error Selenium en {lang}: {e}")
        finally:
            if driver:
                driver.quit()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Selenium - Total: {len(df)} documentos")
    return df


def load_investigacion_bid_unified(start_date_str, end_date_str):
    """
    UNIFICADOR: Prueba cloudscraper primero, si falla usa Selenium
    """
    print("="*50)
    print("🔍 Iniciando extracción BID con estrategia unificada")
    print("="*50)
    
    # Intentar primero con cloudscraper
    try:
        print("\n🚀 Estrategia 1: Cloudscraper")
        df = load_investigacion_bid_cloudscraper(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Cloudscraper exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Cloudscraper no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Cloudscraper falló: {e}")
    
    # Fallback a Selenium
    print("\n🚀 Estrategia 2: Selenium con delay largo")
    try:
        df = load_investigacion_bid_selenium_fallback(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Selenium exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Selenium no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Selenium falló: {e}")
    
    print("\n❌ Ambas estrategias fallaron para BID")
    return pd.DataFrame()

# ==========================================
# EXPORTACIÓN A WORD
# ==========================================

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    b = docx.oxml.shared.OxmlElement('w:b')
    rPr.append(b)

    for s in ['w:sz', 'w:szCs']:
        sz = docx.oxml.shared.OxmlElement(s)
        sz.set(docx.oxml.shared.qn('w:val'), '28')
        rPr.append(sz)

    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri')
    rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_word(df, title="Boletín Mensual", subtitle=""):
    doc = Document()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(df.columns)-1)
    table.style = 'Table Grid'

    cols = [c for c in df.columns if c != 'Link']

    for idx, name in enumerate(cols):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            p = cells[i].paragraphs[0]
            if col == 'Nombre de Documento':
                add_hyperlink(p, str(row[col]), str(row['Link']))
            else:
                run = p.add_run(str(row[col]))
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True

    if 'Tipo de Documento' in df.columns and 'Organismo' in df.columns:
        col_tipo = cols.index('Tipo de Documento')
        col_org = cols.index('Organismo')

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            org_val = df.iloc[start_row - 1]['Organismo']
            end_row = start_row

            if cat_val == "Discursos":
                table.cell(start_row, col_org).text = ""
                while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == "Discursos":
                    table.cell(end_row + 1, col_org).text = ""
                    end_row += 1

                if end_row > start_row:
                    target_cell = table.cell(start_row, col_org)
                    target_cell.merge(table.cell(end_row, col_org))

                start_row = end_row + 1
                continue

            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val and df.iloc[end_row]['Organismo'] == org_val:
                table.cell(end_row + 1, col_org).text = ""
                end_row += 1

            if end_row > start_row:
                target_cell = table.cell(start_row, col_org)
                target_cell.merge(table.cell(end_row, col_org))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            start_row = end_row + 1

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            end_row = start_row

            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val:
                table.cell(end_row + 1, col_tipo).text = ""
                end_row += 1

            if end_row > start_row:
                target_cell = table.cell(start_row, col_tipo)
                target_cell.merge(table.cell(end_row, col_tipo))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            start_row = end_row + 1

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ==========================================
# INTERFAZ DE USUARIO Y MAIN
# ==========================================
try:
    st.sidebar.image("logo_banxico.png", use_container_width=True)
except:
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")
modo_app = st.sidebar.radio(
    "", ["Boletín", "Categorías", "Carga Manual"], key="menu_principal")
st.sidebar.markdown("---")

anios_str = ["2026", "2025", "2024", "2023", "2022"]
meses_dict = {
    "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
    "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
}

# --- LISTAS DINÁMICAS DE ORGANISMOS ---
orgs_discursos = ["BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", "BoC (Canadá)", "BoE (Inglaterra)", "BoJ (Japón)", "BPI", "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"]
orgs_reportes = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
orgs_pub_inst = ["BM", "BPI", "CEF", "CEMLA", "FMI", "FMI (Mission Concluding)", "F&D", "G20", "OCDE", "OEI"]
orgs_investigacion = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]

if modo_app == "Boletín":
    st.title("Generador de Boletín Mensual")
    st.markdown(
        "Extrae y unifica documentos de todas las categorías y organismos por mes.")
    st.markdown("---")

    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])


    if st.button("📄 Generar Boletín Mensual", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

            all_dfs = []
            prog = st.progress(0)
            txt = st.empty()

            total_pasos = len(orgs_discursos) + len(orgs_reportes) + \
                len(orgs_pub_inst) + len(orgs_investigacion)
            paso_actual = 0

            # 1. BARRIDO DE DISCURSOS
            for org in orgs_discursos:
                txt.text(f"Procesando Discursos: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI":
                        df = load_data_bis()
                    elif org == "ECB (Europa)":
                        df = load_data_ecb(sd, ed)
                    elif org == "FMI":
                        df = load_discursos_fmi(sd, ed)
                    elif org == "BBk (Alemania)":
                        df = load_data_bbk(sd, ed)
                    elif org == "Fed (Estados Unidos)":
                        df = load_data_fed(a_num)
                    elif org == "BdF (Francia)":
                        df = load_data_bdf(sd, ed)
                    elif org == "BM":
                        df = load_data_bm(sd, ed)
                    elif org == "BoC (Canadá)":
                        df = load_data_boc(sd, ed)
                    elif org == "BoJ (Japón)":
                        df = load_data_boj(sd, ed)
                    elif org == "BoJ (Japón)": df = load_data_boj(sd, ed)
                    elif org == "BoE (Inglaterra)": df = load_discursos_boe(sd, ed)
                    elif org == "CEF": df = load_data_cef(sd, ed)
                    elif org == "CEF":
                        df = load_data_cef(sd, ed)
                    elif org == "PBoC (China)":
                        df = load_data_pboc(sd, ed)
                    if org == "BPI":
                        df = load_data_bis()
                    elif org == "ECB (Europa)":
                        df = load_data_ecb(sd, ed)
                    elif org == "FMI":
                        df = load_discursos_fmi(sd, ed)
                    elif org == "BBk (Alemania)":
                        df = load_data_bbk(sd, ed)
                    elif org == "BdE (España)":
                        df = load_data_bde(sd, ed)
                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Discursos"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 2. BARRIDO DE REPORTES
            for org in orgs_reportes:
                txt.text(f"Procesando Reportes: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID":
                        df = load_reportes_bid_en(sd, ed)
                    elif org == "BM":
                        df = load_reportes_bm(sd, ed) 
                    elif org == "BPI":
                        df = load_reportes_bpi(sd, ed)
                    elif org == "CEF":
                        df = load_reportes_cef(sd, ed)
                    elif org == "OCDE":
                        df = load_reportes_ocde(sd, ed)
                    elif org == "FEM": 
                        df = load_reportes_fem(sd, ed)
                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Reportes"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES
            for org in orgs_pub_inst:
                txt.text(f"Procesando Pub. Institucionales: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI":
                        df = load_pub_inst_bpi(sd, ed)
                    elif org == "CEF":
                        df = load_pub_inst_cef(sd, ed)
                    elif org == "BM":
                        df = load_pub_inst_bm(sd, ed)
                    elif org == "OEI": 
                        df = load_pub_inst_oei(sd, ed)
                    elif org == "F&D":  
                        df = load_pub_inst_fandd(sd, ed)
                    elif org == "FMI":
                        # 1. SSG - JSON Estático (WEO, Fiscal Monitor)
                        df_flagships = load_pub_inst_fmi(sd, ed)

                        # 2. SSG - JSON Estático (Comunicados)
                        df_prs = load_press_releases_fmi(sd, ed)

                        # 3. CSR API - Coveo (Country Reports)
                        df_crs = load_country_reports_fmi(
                            sd, ed)  # <-- LA NUEVA API

                        # Unión
                        dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.sort_values("Date", ascending=False)

                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Publicaciones Institucionales"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 4. BARRIDO DE INVESTIGACIÓN
            for org in orgs_investigacion:
                txt.text(f"Procesando Investigación: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID": 
                        df = load_investigacion_bid_unified(sd, ed)
                    elif org == "BPI": 
                        df = load_investigacion_bpi(sd, ed)
                    elif org == "BM": 
                        df = load_investigacion_bm(sd, ed)
                    elif org == "CEMLA":
                        df = load_investigacion_cemla(sd, ed)
                    elif org == "FMI": 
                        df_blogs = pd.DataFrame()
                        df_wp = pd.DataFrame()
                        try:
                            df_blogs = load_investigacion_fmi(sd, ed)
                        except: pass
                        try:
                            df_wp = load_working_papers_fmi(sd, ed)
                        except: pass
                        # Unimos Blogs y Working Papers
                        dfs_a_unir = [d for d in [df_blogs, df_wp] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.drop_duplicates(subset=['Link'])
                            df = df.sort_values("Date", ascending=False)
                except Exception as e: pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Investigación"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # --- INYECCIÓN DE TEXTO MANUAL ---
            if 'df_extra' in st.session_state and not st.session_state.df_extra.empty:
                all_dfs.append(st.session_state.df_extra)
                txt.text("Inyectando reportes manuales...")

            txt.empty()
            prog.empty()

            # --- CONSOLIDACIÓN FINAL ---
            if all_dfs:
                f_df = pd.concat(all_dfs, ignore_index=True)

                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                df_rep = f_df[f_df['Categoría'] == "Reportes"].copy()
                df_pub = f_df[f_df['Categoría'] ==
                              "Publicaciones Institucionales"].copy()
                df_inv = f_df[f_df['Categoría'] == "Investigación"].copy()
                df_disc = f_df[f_df['Categoría'] == "Discursos"].copy()

                if not df_rep.empty:
                    df_rep = df_rep.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_pub.empty:
                    df_pub = df_pub.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_inv.empty:
                    df_inv = df_inv.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_disc.empty:
                    df_disc = df_disc.sort_values(
                        by=["Title"], ascending=[True])

                f_df_word = pd.concat(
                    [df_rep, df_pub, df_inv, df_disc], ignore_index=True)
                f_df_word = f_df_word[['Categoría',
                                       'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(
                    columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})

                st.success(
                    f"Se consolidaron **{len(f_df)}** documentos en total.")
                word = generate_word(f_df_word, subtitle=", ".join(
                    m_sel) + " " + ", ".join(a_sel))
                
                # Botón de Descarga
                st.download_button("📄 Descargar Boletín en Word",
                                   word, f"Boletin_{'_'.join(m_sel)}.docx")

                # Limpiar cache manual después de generar el Word para no duplicar en el futuro
                if 'df_extra' in st.session_state:
                    del st.session_state.df_extra

                # --- PREPARACIÓN PARA LA VISTA PREVIA ---
                disp = f_df.copy()
                disp = disp.sort_values(
                    by="Date", ascending=False)  # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime('%d/%m/%Y')
                disp["Nombre de Documento"] = disp.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})

                st.markdown(disp[["Fecha", "Tipo de Documento", "Organismo", "Nombre de Documento"]].to_markdown(
                    index=False), unsafe_allow_html=True)
            else:
                st.warning(
                    "No se encontraron documentos para los criterios seleccionados.")

elif modo_app == "Categorías":
    st.title("Documentos de Organismos Internacionales")
    tipo_doc = st.sidebar.selectbox("Tipo de Documento", [
                                    "Discursos", "Reportes", "Investigación", "Publicaciones Institucionales"])

    # Construcción segura de las listas de interfaz
    if tipo_doc == "Discursos":
        orgs_list = ["Todos"] + sorted(orgs_discursos)
    elif tipo_doc == "Reportes":
        orgs_list = ["Todos"] + sorted(orgs_reportes)
    elif tipo_doc == "Investigación":
        orgs_list = ["Todos"] + sorted(orgs_investigacion)
    elif tipo_doc == "Publicaciones Institucionales":
        orgs_list = ["Todos"] + sorted(orgs_pub_inst)
    else:
        orgs_list = ["Todos"] + sorted(
            list(set(orgs_discursos + orgs_reportes + orgs_investigacion + orgs_pub_inst)))

    organismo_seleccionado = st.sidebar.selectbox("Organismo", orgs_list)

    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])

    if st.button("🔍 Buscar", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

            target_orgs = orgs_list[1:] if organismo_seleccionado == "Todos" else [
                organismo_seleccionado]
            dfs_comb = []
            progreso = st.progress(0)
            txt = st.empty()

            for i, o in enumerate(target_orgs):
                txt.text(f"Extrayendo: {o}...")
                df = pd.DataFrame()
                try:
                    # --- LÓGICA DE EXTRACCIÓN POR TIPO ---
                    if tipo_doc == "Discursos":
                        if o == "BPI":
                            df = load_data_bis()
                        elif o == "ECB (Europa)":
                            df = load_data_ecb(sd, ed)
                        elif o == "BBk (Alemania)":
                            df = load_data_bbk(sd, ed)
                        elif o == "Fed (Estados Unidos)":
                            df = load_data_fed(a_num)
                        elif o == "BdF (Francia)":
                            df = load_data_bdf(sd, ed)
                        elif o == "BM":
                            df = load_data_bm(sd, ed)
                        elif o == "BoC (Canadá)":
                            df = load_data_boc(sd, ed)
                        elif o == "BoJ (Japón)":
                            df = load_data_boj(sd, ed)
                        elif o == "BoE (Inglaterra)": df = load_discursos_boe(sd, ed)
                        elif o == "CEMLA": 
                            print("🔴🔴🔴 LLAMANDO A CEMLA INVESTIGACIÓN 🔴🔴🔴")
                            df = load_investigacion_cemla(sd, ed)
                            print(f"🔴🔴🔴 RESULTADO CEMLA: {len(df)} documentos 🔴🔴🔴")   
                        elif o == "CEF":
                            df = load_data_cef(sd, ed)
                        elif o == "FMI":
                            df = load_discursos_fmi(sd, ed)
                        elif o == "PBoC (China)":
                            df = load_data_pboc(sd, ed)
                        elif o == "BdE (España)":
                            df = load_data_bde(sd, ed)

                    elif tipo_doc == "Reportes":
                        if o == "BID":
                            dfs_bid = []
                            try:
                                dfs_bid.append(load_reportes_bid(sd, ed))
                            except:
                                pass
                            try:
                                dfs_bid.append(load_reportes_bid_en(sd, ed))
                            except:
                                pass
                            dfs_bid = [d for d in dfs_bid if not d.empty]
                            if dfs_bid:
                                df = pd.concat(dfs_bid, ignore_index=True).drop_duplicates(
                                    subset=['Link'])
                        elif o == "BM":
                            df = load_reportes_bm(sd, ed)
                        elif o == "BPI":
                            df = load_reportes_bpi(sd, ed)
                        elif o == "CEF":
                            df = load_reportes_cef(sd, ed)
                        elif o == "OCDE":
                            df = load_reportes_ocde(sd, ed)
                        elif o == "FEM": df = load_reportes_fem(sd, ed)

                    elif tipo_doc == "Investigación":
                        if o == "BID":
                            df = load_investigacion_bid_unified(sd, ed) 
                        elif o == "BPI":
                            df = load_investigacion_bpi(sd, ed)
                        elif o == "BM":
                            df = load_investigacion_bm(sd, ed)
                        elif o == "BPI":
                            df = load_investigacion_bpi(sd, ed)
                        elif o == "BM":
                            df = load_investigacion_bm(sd, ed)
                            
                        elif o == "FMI": 
                            df_blogs, df_wp = pd.DataFrame(), pd.DataFrame()
                            try: df_blogs = load_investigacion_fmi(sd, ed)
                            except: pass
                            try: df_wp = load_working_papers_fmi(sd, ed)
                            except: pass
                            
                            dfs_fmi = [d for d in [df_blogs, df_wp] if not d.empty]
                            if dfs_fmi:
                                df = pd.concat(dfs_fmi, ignore_index=True).drop_duplicates(subset=['Link']).sort_values("Date", ascending=False)
                    elif tipo_doc == "Publicaciones Institucionales":
                        if o == "BPI":
                            df = load_pub_inst_bpi(sd, ed)
                        elif o == "CEF":
                            df = load_pub_inst_cef(sd, ed)
                        elif o == "OEI": df = load_pub_inst_oei(sd, ed)
                        elif o == "BM":
                            df = load_pub_inst_bm(sd, ed)
                        elif o == "FMI":
                            df_flagships = load_pub_inst_fmi(sd, ed)
                            df_prs = load_press_releases_fmi(sd, ed)
                            df_crs = load_country_reports_fmi(sd, ed)
                            dfs_a_unir = [d for d in [
                                df_flagships, df_prs, df_crs] if not d.empty]
                            if dfs_a_unir:
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.sort_values("Date", ascending=False)

                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = o
                        dfs_comb.append(df_f)
                progreso.progress((i+1)/len(target_orgs))

            txt.empty()
            progreso.empty()

            if dfs_comb:
                f_df = pd.concat(dfs_comb, ignore_index=True)
                f_df['Categoría'] = tipo_doc

                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                if tipo_doc == "Discursos":
                    f_df_word = f_df.sort_values(
                        by=["Title"], ascending=[True])
                else:
                    f_df_word = f_df.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])

                f_df_word = f_df_word[['Categoría',
                                       'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(
                    columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})

                st.success(f"Se encontraron **{len(f_df)}** documentos.")
                word_file = generate_word(
                    f_df_word, title=f"Explorador - {tipo_doc}")
                st.download_button(
                    "📄 Descargar en Word", data=word_file, file_name=f"Explorador_{tipo_doc}.docx")

                # --- PREPARACIÓN PARA LA VISTA PREVIA ---
                disp = f_df.copy()
                disp = disp.sort_values(
                    by="Date", ascending=False)  # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime(
                    '%d/%m/%Y')  # Formatear fecha
                disp["Nombre de Documento"] = disp.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})

                if organismo_seleccionado == "Todos":
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Organismo", "Nombre de Documento"]
                else:
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Nombre de Documento"]

                st.markdown(disp[cols_vis].to_markdown(index=False), unsafe_allow_html=True)
            else:
                st.warning(
                    "No se encontraron documentos para las fechas seleccionadas.")

elif modo_app == "Carga Manual":
    st.title("🛠️ Centro de Carga Manual")
    st.markdown("Pega el texto de las páginas que fallan. Previsualiza, valida y une todo en un solo documento.")
    
    if 'cargas_validadas' not in st.session_state:
        st.session_state.cargas_validadas = {
            "OCDE (Reportes)": pd.DataFrame(),
            "OCDE (Pub. Institucionales)": pd.DataFrame(),
            "OCDE (Investigación)": pd.DataFrame()
        }

    st.subheader("Estado de Carga")
    cols_estado = st.columns(3)
    claves_cajas = list(st.session_state.cargas_validadas.keys())
    
    for i, clave in enumerate(claves_cajas):
        estado = "✅ Listo" if not st.session_state.cargas_validadas[clave].empty else "❌ Pendiente"
        cols_estado[i].info(f"**{clave}**\n\n{estado}")

    st.markdown("---")
    
    c1, c2 = st.columns(2)
    mes_manual = c1.selectbox("Mes objetivo a filtrar:", [1,2,3,4,5,6,7,8,9,10,11,12], index=datetime.datetime.now().month-1, format_func=lambda x: calendar.month_name[x].capitalize())
    año_manual = c2.number_input("Año objetivo a filtrar:", min_value=2020, max_value=2030, value=datetime.datetime.now().year)

    st.markdown("---")
    st.subheader("Cajas de Extracción")

    def crear_caja_manual(titulo_caja, categoria_doc, organismo_nombre, url_fuente=None):
        with st.expander(f"📥 Cargar: {titulo_caja}", expanded=True):
            
            if url_fuente:
                st.markdown(f"👉 **[Haz clic aquí para abrir la página oficial de {titulo_caja}]({url_fuente})**")
                
            texto = st.text_area(f"Copia el texto de la página y pégalo aquí (Ctrl+A, Ctrl+C, Ctrl+V):", height=150, key=f"txt_{titulo_caja}")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            
            if col_btn1.button(f"🔍 Previsualizar {titulo_caja}", key=f"btn_prev_{titulo_caja}"):
                if texto:
                    with st.spinner("Procesando y buscando links..."):
                        df_bruto = procesar_texto_pegado(texto, organismo_nombre)
                            
                        if not df_bruto.empty:
                            df_filtrado = df_bruto[
                                (df_bruto['Date'].dt.month == mes_manual) & 
                                (df_bruto['Date'].dt.year == año_manual)
                            ].copy()
                            
                            if not df_filtrado.empty:
                                for idx in df_filtrado.index:
                                    t = df_filtrado.loc[idx, "Title"]
                                    df_filtrado.loc[idx, "Link"] = buscar_link_inteligente(t, organismo_nombre)
                                
                                df_filtrado['Categoría'] = categoria_doc
                                st.session_state[f"temp_{titulo_caja}"] = df_filtrado
                                
                                st.success(f"Se encontraron {len(df_filtrado)} documentos de {mes_manual} {año_manual}.")
                                st.dataframe(df_filtrado, use_container_width=True)
                            else:
                                st.warning("No hay coincidencias con el mes y año seleccionados.")
                else:
                    st.error("Pega el texto primero.")
            
            if col_btn2.button(f"➕ Agregar a Descarga Final", type="primary", key=f"btn_add_{titulo_caja}"):
                if f"temp_{titulo_caja}" in st.session_state and not st.session_state[f"temp_{titulo_caja}"].empty:
                    st.session_state.cargas_validadas[titulo_caja] = st.session_state[f"temp_{titulo_caja}"]
                    st.success(f"¡{titulo_caja} guardado en memoria! ✅")
                    time.sleep(1)
                    st.rerun() 
                else:
                    st.error("Primero debes Previsualizar y obtener resultados.")

    link_ocde_rep = "https://www.oecd.org/en/search/publications.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Freports%2Coecd-languages%3Aen&minPublicationYear=2026&maxPublicationYear=2026"
    link_ocde_pub = "https://www.oecd.org/en/search.html?orderBy=mostRecent&page=0&facetTags=oecd-policy-subissues%3Apsi114%2Coecd-languages%3Aen"
    link_ocde_inv = "https://www.oecd.org/en/publications/reports.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Fworking-papers%2Coecd-languages%3Aen"
    
    crear_caja_manual("OCDE (Reportes)", "Reportes", "OCDE", link_ocde_rep)
    crear_caja_manual("OCDE (Pub. Institucionales)", "Publicaciones Institucionales", "OCDE", link_ocde_pub)
    crear_caja_manual("OCDE (Investigación)", "Investigación", "OCDE", link_ocde_inv)

    st.markdown("---")
    st.subheader("Exportación Final")
    
    tablas_listas = [df for df in st.session_state.cargas_validadas.values() if not df.empty]
    
    if tablas_listas:
        df_maestro = pd.concat(tablas_listas, ignore_index=True)
        num_cat = len(tablas_listas)
        st.info(f"Tienes **{num_cat}/3** categorías listas, sumando un total de **{len(df_maestro)}** documentos para exportar.")
        
        df_word_manual = df_maestro[['Categoría', 'Organismo', 'Title', 'Link']].copy()
        df_word_manual = df_word_manual.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
        
        word_manual = generate_word(df_word_manual, title="Boletín - Carga Manual", subtitle=f"Mes: {mes_manual} | Año: {año_manual}")
        
        c_down, c_clear = st.columns(2)
        with c_down:
            st.download_button(
                label=f"📄 Descargar Word ({num_cat}/3 Listas)", 
                data=word_manual, 
                file_name=f"Carga_Manual_{mes_manual}_{año_manual}.docx"
            )
        with c_clear:
            if st.button("🗑️ Reiniciar todo el módulo"):
                for clave in st.session_state.cargas_validadas.keys():
                    st.session_state.cargas_validadas[clave] = pd.DataFrame()
                st.rerun()
    else:
        st.warning("Aún no has agregado ninguna carga a la descarga final. Agrega al menos 1 para habilitar el botón de descarga.")
